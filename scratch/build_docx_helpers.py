import os
import glob
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{top}" w:type="dxa"/><w:bottom w:w="{bottom}" w:type="dxa"/><w:left w:w="{left}" w:type="dxa"/><w:right w:w="{right}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_heading_1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(16)
    run.font.bold = True
    run.font.color.rgb = RGBColor(15, 23, 42) # Dark Slate #0f172a
    return p

def add_heading_2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(13)
    run.font.bold = True
    run.font.color.rgb = RGBColor(30, 58, 138) # Deep Blue #1e3a8a
    return p

def add_heading_3(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(11)
    run.font.bold = True
    run.font.color.rgb = RGBColor(99, 102, 241) # Indigo #6366f1
    return p

def add_paragraph(doc, text, bold_prefix="", italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Arial'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.italic = italic
    run.font.color.rgb = RGBColor(51, 65, 85) # Slate #334155
    return p

def add_bullet(doc, text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    if bold_prefix:
        r_pre = p.add_run(bold_prefix)
        r_pre.font.name = 'Arial'
        r_pre.font.size = Pt(10.5)
        r_pre.font.bold = True
        r_pre.font.color.rgb = RGBColor(15, 23, 42)
    
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(51, 65, 85)
    return p

def add_image_placeholder(doc, title, description, instructions=""):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = False
    
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, "EFF6FF") # Light blue tint
    set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
    
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    
    r_tag = p.add_run(f"📌 {title.upper()}\n")
    r_tag.font.name = 'Arial'
    r_tag.font.size = Pt(11)
    r_tag.font.bold = True
    r_tag.font.color.rgb = RGBColor(29, 78, 216) # Royal Blue #1d4ed8
    
    p_desc = cell.add_paragraph()
    p_desc.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_desc.paragraph_format.space_before = Pt(4)
    p_desc.paragraph_format.space_after = Pt(4)
    r_desc = p_desc.add_run(f"📝 Keterangan Gambar:\n{description}")
    r_desc.font.name = 'Arial'
    r_desc.font.size = Pt(9.5)
    r_desc.font.color.rgb = RGBColor(30, 41, 59)
    
    if instructions:
        p_inst = cell.add_paragraph()
        p_inst.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_inst.paragraph_format.space_before = Pt(2)
        p_inst.paragraph_format.space_after = Pt(6)
        r_inst = p_inst.add_run(f"💡 Panduan Menempelkan Gambar:\n{instructions}")
        r_inst.font.name = 'Arial'
        r_inst.font.size = Pt(9)
        r_inst.font.italic = True
        r_inst.font.color.rgb = RGBColor(71, 85, 105)
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

def add_code_block(doc, code_text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.rows[0].cells[0]
    cell.width = Inches(6.5)
    set_cell_background(cell, "0F172A") # Dark background
    set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(56, 189, 248) # Sky blue #38bdf8
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

print("Helper script template created!")
