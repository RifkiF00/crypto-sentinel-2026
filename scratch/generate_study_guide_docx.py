"""
Script untuk menghasilkan berkas Microsoft Word (.docx) resmi:
docs/PANDUAN_STUDI_LITERATUR_DAN_DEFENSE_TIM.docx
"""
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(level=level)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    if level == 1:
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        r.font.size = Pt(15)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor(2, 132, 199)
    elif level == 3:
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        r.font.size = Pt(10.5)
        r.bold = True
        r.font.color.rgb = RGBColor(51, 65, 85)
    return h

def add_callout(doc, title, text, bg="F8FAFC", border_color="0284C7"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, bg)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if title:
        r_t = p.add_run(title + "\n")
        r_t.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.color.rgb = RGBColor(15, 23, 42)
    r_body = p.add_run(text)
    r_body.font.size = Pt(9)
    r_body.font.color.rgb = RGBColor(51, 65, 85)

def build_docx():
    doc = docx.Document()
    
    # Page Setup
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Title Hero
    t_hero = doc.add_table(rows=1, cols=1)
    t_hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_hero = t_hero.cell(0, 0)
    c_hero.width = Inches(6.8)
    set_cell_background(c_hero, "09132E")
    set_cell_margins(c_hero, top=200, bottom=200, left=180, right=180)
    
    p_b = c_hero.paragraphs[0]
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_b = p_b.add_run("PIDI DIGDAYA 2026 · TEAM EXPRESSO S1251\n")
    r_b.font.size = Pt(9)
    r_b.bold = True
    r_b.font.color.rgb = RGBColor(217, 119, 6)
    
    r_t = p_b.add_run("PANDUAN STUDI LITERATUR & Q&A DEFENSE\n")
    r_t.font.size = Pt(18)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(255, 255, 255)
    
    r_s = p_b.add_run("Master Buku Saku Penguasaan Teori, Arsitektur, Rumus, dan Naskah Jawaban Juri\n")
    r_s.font.size = Pt(10)
    r_s.font.color.rgb = RGBColor(148, 163, 184)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. Kerangka Pikir
    add_heading_styled(doc, "1. Kerangka Pikir 'Vibe Coder to Enterprise Defender'", level=1)
    add_callout(doc, "💡 Mengapa Tim Vibe Coder Bisa Menang di Hadapan Juri?",
                "Vibe coder yang hebat bukan yang berpura-pura menulis setiap baris kode dari nol, melainkan yang memahami arsitektur sistem secara holistik, memahami mengapa setiap keputusan desain diambil, dan mampu menjelaskan logika sistem dengan presisi standar perbankan.\n\n"
                "Teknik Menjawab Piramida Top-Down:\n"
                "1. Level 1 (Direct Answer - 5 Detik): Jawab langsung inti pertanyaan ('Ya/Tidak, sistem kami menerapkan pola...').\n"
                "2. Level 2 (Mechanism & Code - 15 Detik): Sebutkan modul dan caranya ('Di file rule_engine.py kami menghitung Haversine...').\n"
                "3. Level 3 (Business & Regulation - 10 Detik): Tutup dengan dampaknya ('Menjamin kepatuhan POJK No. 8/2023 dan latensi <25ms').",
                bg="F0F9FF", border_color="0284C7")

    # 2. Rifki
    add_heading_styled(doc, "Bagian 1: Rifki Firmansyah — AI Architecture, GNN & Benchmark", level=1)
    p_rif = doc.add_paragraph()
    p_rif.add_run("Role: Team Lead, AI Architect & Product Strategist\nOwnership: main.py, rule_engine.py, gnn_scorer.py, ml_model.joblib, eval_official_benchmark.py").italic = True

    add_heading_styled(doc, "A. Teori Ilmiah & Studi Literatur AI", level=2)
    ai_lit = [
        ("GraphSAGE (Hamilton et al., NIPS 2017)", "Deep learning induktif untuk representasi graf dinamis. Graf dibangun dari 562.239 node rekening dan 308.213 edge relasi transfer PaySim. Fitur node 8-dimensi menghasilkan embedding representatif 32-dimensi via training GPU T4 (Adam lr=0.005, Loss konvergen 0.0004 di Epoch 15)."),
        ("Random Forest (Breiman, 2001) + SMOTE", "Ensemble 100 pohon keputusan memproses 29 fitur tabular. SMOTE (Chawla et al., 2002) menyeimbangkan kelas fraud dari 0.53% menjadi 50:50 pada fase fitting."),
        ("GNNExplainer (Ying et al., NeurIPS 2019)", "Memaksimalkan Mutual Information: max MI(Y, (G_s, F)) = H(Y) - H(Y | G=G_s, X=X*F) untuk mengisolasi subgraf sindikat mule 3-hop."),
        ("Formula Fusi Hibrida", "hybrid = 0.6 * gnn_score + 0.4 * rule_score  ;  final_score = max(hybrid, rule). Fungsi max() menjamin safety floor policy.")
    ]
    for k, v in ai_lit:
        p = doc.add_paragraph(style='List Bullet')
        r_k = p.add_run(k + ": ")
        r_k.bold = True
        p.add_run(v)

    add_heading_styled(doc, "B. Official Benchmark Kinerja AI (308.213 Data PaySim)", level=2)
    bench_data = [
        ("Akurasi Keseluruhan", "99.98%"),
        ("Presisi (Precision)", "99.94%"),
        ("Recall (Sensitivitas Fraud)", "99.88% (1.641 / 1.643 fraud tertangkap)"),
        ("F1-Score", "99.91%"),
        ("ROC-AUC", "0.9997"),
        ("False Positive Rate (FPR)", "0.002% (Hanya 1 dari 60.000 transaksi legitimate)"),
        ("False Negative Rate (FNR)", "0.122% (Hanya 2 fraud terlewat dari 1.643 kasus)"),
        ("Inference Latency Mean", "5.67 ms (p95: 9.05 ms, p99: 12.23 ms pada CPU lokal)")
    ]
    t_b = doc.add_table(rows=len(bench_data)+1, cols=2)
    set_table_borders(t_b)
    t_b.cell(0, 0).paragraphs[0].add_run("Metrik Evaluasi").bold = True
    t_b.cell(0, 1).paragraphs[0].add_run("Nilai Resmi Terverifikasi").bold = True
    set_cell_background(t_b.cell(0, 0), "0F172A")
    set_cell_background(t_b.cell(0, 1), "0F172A")
    t_b.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    t_b.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, (m, val) in enumerate(bench_data):
        c0 = t_b.cell(i+1, 0)
        c1 = t_b.cell(i+1, 1)
        c0.paragraphs[0].add_run(m).bold = True
        c1.paragraphs[0].add_run(val)
        bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)

    add_callout(doc, "🎤 Naskah Q&A Defense Rifki (Jawaban Wajib Hafal):",
                "Q: Apakah sistem melakukan message-passing GNN online di setiap request? Bagaimana latensinya bisa 5.67 ms?\n"
                "A: 'Secara arsitektur enterprise, menjalankan deep message-passing online pada setiap transaksi perbankan akan membebani latensi. Oleh karena itu, GraphSAGE dilatih secara offline pada 562 ribu entitas untuk menghasilkan embedding 32-dimensi. Di runtime API, kami melakukan fast embedding lookup yang digabungkan dengan Random Forest dan Rule Engine pada CPU ringan. Inilah mengapa latensi rata-rata kami 5.67 milidetik tanpa kehilangan keunggulan deteksi graf relasional.'",
                bg="FFFBEB", border_color="D97706")

    # 3. Billy
    add_heading_styled(doc, "Bagian 2: Billy Jonathan — Cyber Security, SNAP BI & CBS Gateway", level=1)
    p_bil = doc.add_paragraph()
    p_bil.add_run("Role: Cyber Security & Core Banking Developer\nOwnership: transfers.py, db_models.py, otentikasi SNAP BI HMAC-SHA256, CBS Gateway").italic = True

    sec_pts = [
        ("SNAP BI (PADG No. 23/15/PADG/2021)", "Standar Nasional Open API Pembayaran Bank Indonesia. Menggunakan HMAC-SHA256 signature, X-TIMESTAMP drift limit 5 menit, dan X-PARTNER-ID (KNG-BANK-001 / BJB-BANK-002)."),
        ("Smart Circuit Breaker Pattern (Nygard, 2007)", "Pencegatan transaksi pre-authorization di middleware sebelum debit ledger dieksekusi di Core Banking. Jika skor >= 85, transaksi ditolak dalam <25ms dan saldo nasabah aman 100%."),
        ("Anatomi Serangan Smurfing & Drain-to-Zero", "Smurfing: memecah Rp 500 juta menjadi pecahan Rp 4,9 juta ke 100 rekening mule. Drain-to-zero: menguras saldo hingga nol di jam 02:00 WIB ke bursa kripto.")
    ]
    for k, v in sec_pts:
        p = doc.add_paragraph(style='List Bullet')
        r_k = p.add_run(k + ": ")
        r_k.bold = True
        p.add_run(v)

    add_callout(doc, "🎤 Naskah Q&A Defense Billy (Jawaban Wajib Hafal):",
                "Q: Bagaimana sistem mencegah pemalsuan request (Man-in-the-Middle) di mobile banking?\n"
                "A: 'Setiap payload transaksi wajib menyertakan header SNAP BI resmi: X-TIMESTAMP, X-PARTNER-ID, dan X-SIGNATURE. Di gateway transfers.py, sistem menghitung ulang hash HMAC-SHA256 dari body payload. Jika ada perubahan nominal 1 rupiah saja, signature mismatch seketika dan request ditolak dengan HTTP 401 Unauthorized.'",
                bg="EFF6FF", border_color="2563EB")

    # 4. Aam
    add_heading_styled(doc, "Bagian 3: Aam Setiana — Frontend UI/UX, Forensic Dashboard & Privacy PDP", level=1)
    p_aam = doc.add_paragraph()
    p_aam.add_run("Role: Frontend Engineer & Product Analyst\nOwnership: App.jsx, PlatformViews.jsx, PageViews.jsx, masking.js, AuthContext.jsx").italic = True

    fe_pts = [
        ("UU PDP No. 27/2022 (Pasal 35 & 36)", "Mandat enkripsi dan pseudonimisasi data. Di masking.js, nama disamarkan B*** S*******, rekening ****6666, dan NIK 3208**********02 secara default. Pembukaan sensor dicatat di audit log."),
        ("3-Tier RBAC (NIST Standard)", "Pemisahan peran: Analis AML (triage & LTKM preview), Compliance Officer / MLRO (otorisasi blokir & threshold POJK), Pengawas OJK (read-only audit APOLO)."),
        ("Live Customer 360 Database", "Memanggil endpoint GET /api/v1/bri/account/{id} untuk menarik data CRA nasabah aktual (PEP status, CDD/EDD, mule probability) langsung dari NeonDB PostgreSQL.")
    ]
    for k, v in fe_pts:
        p = doc.add_paragraph(style='List Bullet')
        r_k = p.add_run(k + ": ")
        r_k.bold = True
        p.add_run(v)

    add_callout(doc, "🎤 Naskah Q&A Defense Aam (Jawaban Wajib Hafal):",
                "Q: Bagaimana dashboard Anda melindungi privasi data nasabah perbankan?\n"
                "A: 'Kami menerapkan Privacy by Design sesuai UU PDP No. 27/2022. Di masking.js dan seluruh platform dashboard, sensor privasi aktif by default. Nama nasabah, NIK, dan nomor rekening otomatis tersamar. Pembukaan data hanya dapat dilakukan role berwenang dan tercatat permanen di audit trail backend.'",
                bg="FDF4FF", border_color="C026D3")

    # 5. Desta
    add_heading_styled(doc, "Bagian 4: Desta Erlangga — Backend Systems, PPATK goAML & Regulasi OJK", level=1)
    p_des = doc.add_paragraph()
    p_des.add_run("Role: Backend & Integration Engineer\nOwnership: db_models.py, NeonDB PostgreSQL, generator draf LTKM goAML di main.py, BANK_INTEGRATION_KIT.md").italic = True

    be_pts = [
        ("POJK No. 8/2023 (Strategi Anti-Fraud)", "Mandat 4 pilar SAF (Pencegahan, Deteksi, Investigasi/Pelaporan, Evaluasi). Penahanan transaksi dan pelaporan audit berkala."),
        ("UU TPPU No. 8/2010 & goAML XML", "Kewajiban pelaporan LTKM / STR ke PPATK. Generator di main.py mengompilasi draf resmi format goAML dalam 3 detik siap verifikasi Compliance Officer."),
        ("Bank Integration Kit (BANK_INTEGRATION_KIT.md)", "Pola integrasi non-intrusif: Mode A (Pre-Auth API Gateway) & Mode B (Post-Transaction CDC Database).")
    ]
    for k, v in be_pts:
        p = doc.add_paragraph(style='List Bullet')
        r_k = p.add_run(k + ": ")
        r_k.bold = True
        p.add_run(v)

    add_callout(doc, "🎤 Naskah Q&A Defense Desta (Jawaban Wajib Hafal):",
                "Q: Apakah sistem Anda otomatis submit laporan ke server PPATK tanpa campur tangan manusia?\n"
                "A: 'Sama sekali tidak. Sesuai regulasi UU TPPU No. 8/2010 dan SOP perbankan, Crypto-Sentinel diposisikan sebagai LTKM / STR Preparation and Compliance Preview Tool. Sistem kami mengotomasi draf formulir goAML dalam 3 detik, namun submit final ke portal PPATK tetap berada di bawah otorisasi manual dan kewenangan Pejabat Kepatuhan (MLRO).' ",
                bg="ECFDF5", border_color="059669")

    # 6. Deep-Dive Regulasi & CMS
    add_heading_styled(doc, "Bagian 5: Deep-Dive Regulasi, Kepatuhan Perbankan & Compliance Management System (CMS)", level=1)
    add_callout(doc, "📜 Doktrin Kepatuhan & Anti Tipping-Off (Pasal 12 UU No. 8/2010):",
                "Mengapa pesan penolakan di mobile banking hanya berbunyi: 'Transaksi Tidak Dapat Diproses. Silakan hubungi CS bank Anda'?\n\n"
                "Penjelasan Hukum: Regulasi perbankan MELARANG KERAS membocorkan kepada nasabah atau pelaku bahwa transaksinya sedang diinvestigasi atau dilaporkan ke PPATK (Anti Tipping-Off). Jika nasabah diberitahu akunnya dicurigai pencucian uang, pelaku akan mematikan rekening dan memindahkan sisa dana sebelum aparat bertindak. Pesan penolakan netral melindungi bank dan mematuhi hukum perundang-undangan.",
                bg="FEF2F2", border_color="DC2626")

    add_heading_styled(doc, "A. Matriks 4 Pilar Strategi Anti-Fraud (POJK No. 8/2023)", level=2)
    saf_data = [
        ("Pilar 1: Pencegahan", "Customer Due Diligence (CDD/EDD), CRA Profiling (pep_status, mule_probability), dan Contextual Whitelist instansi resmi."),
        ("Pilar 2: Deteksi", "Rule Engine 13 indikator + Random Forest 29 fitur + GraphSAGE GNN (<25ms Pre-Authorization Circuit Breaker)."),
        ("Pilar 3: Investigasi & Pelaporan", "Forensic Case Management, Subgraf GNNExplainer 3-Hop, dan otomasi draf resmi goAML LTKM PPATK 3 detik."),
        ("Pilar 4: Evaluasi & Monitoring", "Kalibrasi ambang batas risiko POJK, deteksi drift model, dan audit trail persisten PostgreSQL.")
    ]
    t_saf = doc.add_table(rows=len(saf_data)+1, cols=2)
    set_table_borders(t_saf)
    t_saf.cell(0, 0).paragraphs[0].add_run("Pilar SAF POJK 8/2023").bold = True
    t_saf.cell(0, 1).paragraphs[0].add_run("Implementasi Solusi Crypto-Sentinel 2026").bold = True
    set_cell_background(t_saf.cell(0, 0), "0F172A")
    set_cell_background(t_saf.cell(0, 1), "0F172A")
    t_saf.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    t_saf.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, (p_nom, p_imp) in enumerate(saf_data):
        c0 = t_saf.cell(i+1, 0)
        c1 = t_saf.cell(i+1, 1)
        c0.paragraphs[0].add_run(p_nom).bold = True
        c1.paragraphs[0].add_run(p_imp)
        bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)

    add_heading_styled(doc, "B. Siklus Hidup Kasus Compliance Management System (CMS)", level=2)
    cms_steps = [
        ("1. Ingestion & Pre-Score", "Payload transaksi dievaluasi oleh Rule + RF + GNN dalam waktu <25 ms."),
        ("2. Triage & SLA Assignment", "Kasus diklasifikasikan: Severity Critical (Skor >= 85, SLA T+0 / 4 Jam) vs Severity Medium (Skor 60-84, SLA T+1 / 24 Jam)."),
        ("3. Forensic Workbench", "Analis meneliti Customer 360 (NeonDB live), Subgraf GNN 3-hop, dan kontribusi fitur SHAP."),
        ("4. LTKM Preparation", "Sistem mengompilasi draf formulir goAML PPATK format XML/PDF dalam 3 detik."),
        ("5. Maker-Checker Approval", "Analis mengajukan draf -> Pejabat Kepatuhan (MLRO) mengesahkan dan memutuskan blokir permanen / submit regulator."),
        ("6. Immutable Audit Trail", "Seluruh riwayat aksi tersimpan permanen di tabel audit_logs dengan stempel waktu dan identitas aktor.")
    ]
    for cs_num, cs_txt in cms_steps:
        p = doc.add_paragraph(style='List Bullet')
        r_num = p.add_run(cs_num + ": ")
        r_num.bold = True
        p.add_run(cs_txt)

    # 7. Glosarium
    add_heading_styled(doc, "Glosarium Istilah Kunci (Wajib Kuasai)", level=1)
    glo = [
        ("FDS", "Fraud Detection System — Sistem pencegah transaksi penipuan perbankan."),
        ("Mule Account", "Rekening penampung dana kejahatan yang disewa sindikat."),
        ("Smurfing", "Memecah transaksi besar menjadi pecahan kecil di bawah limit audit."),
        ("Drain-to-Zero", "Pengurasan saldo rekening hingga bersisa Rp 0."),
        ("GraphSAGE", "Graph Sample and Aggregate — Deep learning graf induktif skala enterprise."),
        ("GNNExplainer", "Algoritma XAI subgraf pemisah sindikat berdasarkan Mutual Information."),
        ("SHAP", "Shapley Additive exPlanations — Nilai kontribusi fitur tabular terhadap skor."),
        ("LTKM / STR", "Laporan Transaksi Keuangan Mencurigakan — Format resmi goAML PPATK."),
        ("POJK 8/2023", "Regulasi OJK tentang 4 pilar Strategi Anti-Fraud perbankan."),
        ("UU TPPU 8/2010", "UU Pencegahan Pencucian Uang, kewajiban pelaporan LTKM 3 hari kerja, & larangan tipping-off."),
        ("UU PDP 27/2022", "Regulasi Pelindungan Data Pribadi dan kewajiban masking data nasabah."),
        ("Anti Tipping-Off", "Larangan hukum membocorkan kepada nasabah bahwa akunnya sedang diselidiki.")
    ]
    t_g = doc.add_table(rows=len(glo)+1, cols=2)
    set_table_borders(t_g)
    t_g.cell(0, 0).paragraphs[0].add_run("Istilah / Akronim").bold = True
    t_g.cell(0, 1).paragraphs[0].add_run("Arti & Relevansi Perbankan").bold = True
    set_cell_background(t_g.cell(0, 0), "0F172A")
    set_cell_background(t_g.cell(0, 1), "0F172A")
    t_g.cell(0, 0).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    t_g.cell(0, 1).paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
    for i, (k, v) in enumerate(glo):
        c0 = t_g.cell(i+1, 0)
        c1 = t_g.cell(i+1, 1)
        c0.paragraphs[0].add_run(k).bold = True
        c1.paragraphs[0].add_run(v)
        bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
        set_cell_background(c0, bg)
        set_cell_background(c1, bg)

    output_path = r"d:\Crypto-Sentinel 2026\docs\PANDUAN_STUDI_LITERATUR_DAN_DEFENSE_TIM.docx"
    doc.save(output_path)
    print(f"Successfully generated DOCX at: {output_path}")

if __name__ == "__main__":
    build_docx()
