import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_box_borders(table, color="0284C7", sz="8", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_callout(doc, title, text, bg="F8FAFC", border_color="0284C7"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = tbl.cell(0, 0)
    c.width = Inches(6.9)
    set_cell_background(c, bg)
    set_cell_margins(c, top=120, bottom=120, left=180, right=180)
    
    tblPr = tbl._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="none"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    
    p = c.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if title:
        r_title = p.add_run(title + "\n")
        r_title.bold = True
        r_title.font.size = Pt(9.5)
        r_title.font.color.rgb = RGBColor(15, 23, 42)
    r_text = p.add_run(text)
    r_text.font.size = Pt(9.5)
    r_text.font.italic = True
    r_text.font.color.rgb = RGBColor(51, 65, 85)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(level=level)
    h.paragraph_format.space_before = Pt(14)
    h.paragraph_format.space_after = Pt(4)
    r = h.add_run(text)
    if level == 1:
        r.font.size = Pt(15)
        r.font.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42) # #0F172A
    elif level == 2:
        r.font.size = Pt(12.5)
        r.font.bold = True
        r.font.color.rgb = RGBColor(2, 132, 199) # #0284C7
    elif level == 3:
        r.font.size = Pt(11)
        r.font.bold = True
        r.font.color.rgb = RGBColor(51, 65, 85) # #334155
    return h

def build_pitch_deck_docx(output_path):
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.75)
        section.bottom_margin = Inches(0.75)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    # Base typography
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10)
    normal_style.font.color.rgb = RGBColor(30, 41, 59)

    # 1. HERO HEADER BANNER
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    c = header_table.cell(0, 0)
    c.width = Inches(7.0)
    set_cell_background(c, "09132E")
    set_cell_margins(c, top=200, bottom=200, left=240, right=240)
    
    p0 = c.paragraphs[0]
    p0.paragraph_format.space_before = Pt(0)
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run("PROGRAM PIDI DIGDAYA HACKATHON & INKUBASI 2026 — TIM EXPRESSO S1251")
    r0.font.size = Pt(8.5)
    r0.font.bold = True
    r0.font.color.rgb = RGBColor(56, 189, 248) # Sky blue
    
    p1 = c.add_paragraph()
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after = Pt(2)
    r1 = p1.add_run("CRYPTO-SENTINEL 2026")
    r1.font.size = Pt(18)
    r1.font.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)
    
    p2 = c.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(6)
    r2 = p2.add_run("Master Content Final Pitch Deck, Panduan Eksekusi Presentasi & Supporting Appendix")
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(224, 242, 254)
    
    p3 = c.add_paragraph()
    p3.paragraph_format.space_before = Pt(4)
    p3.paragraph_format.space_after = Pt(0)
    r3 = p3.add_run("Klasifikasi: Live Functional Prototype / Controlled Sandbox Candidate  |  Format Resmi: 12 Slides Standar + 10 Appendix Slides  |  Prinsip: Evidence over Claim · Built over Planned")
    r3.font.size = Pt(8)
    r3.font.color.rgb = RGBColor(148, 163, 184)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 2. CROSS REFERENCES & RUJUKAN TEKNIS
    add_heading_styled(doc, "Dokumen Rujukan Utama & Matriks Referensi Silang (Cross-References)", level=1)
    p_ref = doc.add_paragraph()
    p_ref.add_run("Master Content Pitch Deck ini merupakan materi presentasi eksekutif berbasis bukti (evidence-based) yang didukung langsung oleh dokumen teknis, blueprint arsitektur, dan laporan kemajuan proyek resmi:")
    
    ref_list = [
        ("Project Progress Report (docs/project_progress_report.md)", "Laporan teknis kemajuan proyek 100%, arsitektur dual-mode, evaluasi model AI (Akurasi 99.98%, AUC 0.9993), kepatuhan regulasi POJK 8/2023 & UU PDP, serta hasil uji 8 skenario perbankan."),
        ("Crypto-Sentinel Project Blueprint (crypto-sentinel-blueprint.html)", "Blueprint visual interaktif: arsitektur 4-layer, 15 sub-indikator (4 signal groups), diagram alur data SNAP BI, tahapan pengembangan AI (Stage 1–4), skema database SQLite & PostgreSQL, roadmap, dan profil tim."),
        ("Solution Alignment Notes (docs/solution_alignment_notes.md)", "Notulensi empiris pengujian solusi lapangan bersama calon offtaker PT Bank bjb Tbk dan PT BPR Kuningan (Perseroda)."),
        ("Bank Kuningan Tech Research & Draft LOI (docs/)", "Riset teknologi core banking BPR dan draf Letter of Intent untuk pilot sandbox 3 bulan.")
    ]
    for title, desc in ref_list:
        p_item = doc.add_paragraph(style='List Bullet')
        p_item.paragraph_format.space_before = Pt(1)
        p_item.paragraph_format.space_after = Pt(2)
        r_t = p_item.add_run(title + ": ")
        r_t.bold = True
        r_t.font.color.rgb = RGBColor(2, 132, 199)
        p_item.add_run(desc)

    # Table Quick Jump Map
    add_heading_styled(doc, "Matriks Pemetaan Referensi per Slide Presentasi (Quick Jump Guide)", level=2)
    jump_data = [
        ("Slide 1: Solution at a Glance", "Progress Report: Bab 1.1 Profil Proyek", "Blueprint: Section 00 Overview"),
        ("Slide 2: Problem & Why It Matters", "Progress Report: Bab 1.1 Latar Belakang", "Blueprint: Hero Stats & Scope"),
        ("Slide 3: Validation & Root Cause", "Progress Report: Bab 5.1 Temuan Validasi", "Blueprint: Section 01 Celah Deteksi"),
        ("Slide 4: Solution & Core Use Case", "Progress Report: Bab 1.2 Dual-Mode Engine", "Blueprint: Section 01 Flow Diagram"),
        ("Slide 5: Value Proposition", "Progress Report: Bab 1.3 Keunggulan", "Blueprint: Section 00 & 03 Uniqueness"),
        ("Slide 6: Prototype State", "Progress Report: Bab 2.1 Status Komponen", "Blueprint: Section 02 Stack & 06 DB"),
        ("Slide 7: How Technology Works", "Progress Report: Bab 3.1 AI Hibrida", "Blueprint: Section 03 Indikator & 05 AI"),
        ("Slide 8: Technical Testing", "Progress Report: Bab 4.1 Benchmark AI", "Blueprint: Hero Stats (<50ms, >95% F1)"),
        ("Slide 9: Impact & Effectiveness", "Progress Report: Bab 1.3 Transformasi", "Blueprint: Section 04 Otomasi LTKM"),
        ("Slide 10: Market Validation", "Progress Report: Bab 5.2 Uji 8 Kasus", "Blueprint: Section 07 Roadmap Pasar"),
        ("Slide 11: Adoption & Sustainability", "Progress Report: Bab 6.1 Roadmap Pilot", "Blueprint: Section 07 Roadmap & Biaya"),
        ("Slide 12: Team Readiness", "Progress Report: Bab 1.1 Tim Pengembang", "Blueprint: Section 08 Tim EXPRESSO"),
        ("Appendix A - J (Deep-Dive)", "Progress Report: Bab 2, 3, 4, 5, 6 Lengkap", "Blueprint: Section 01 - 08 Full Spec")
    ]
    t_jump = doc.add_table(rows=len(jump_data)+1, cols=3)
    t_jump.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_jump)
    headers = ["Slide / Bagian", "Rujukan Dokumen Progress Report", "Rujukan Dokumen Blueprint"]
    col_w = [Inches(2.3), Inches(2.5), Inches(2.2)]
    for j, h in enumerate(headers):
        cell = t_jump.cell(0, j)
        cell.width = col_w[j]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(jump_data):
        for j, val in enumerate(row):
            cell = t_jump.cell(i+1, j)
            cell.width = col_w[j]
            bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if j == 0:
                r.bold = True
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 3. ELIGIBILITY CHECK MATRIX
    add_heading_styled(doc, "Matriks Eligibility Check (12 Pertanyaan Lolos Seleksi Halaman 14)", level=1)
    eligibility_data = [
        ("1. Problem", "Apakah masalah benar-benar terjadi didukung bukti nyata?", "Kerugian OJK IASC Rp 9,1T; pembobolan BI-FAST BPD ratusan miliar; data transaksi PaySim 320K.", "LOLOS"),
        ("2. Alignment", "Apakah solusi langsung menjawab problem statement?", "Menghentikan pencucian uang rekening mule dan aliran kripto ilegal sebelum dana keluar bank.", "LOLOS"),
        ("3. Solution", "Dapatkah hubungan problem -> solution dijelaskan tanpa jargon?", "Transfer dicegat di middleware dalam <25ms, dianalisis di dashboard, dan otomatis jadi draf PPATK.", "LOLOS"),
        ("4. Prototype", "Apakah ada bagian inti solusi yang benar-benar sudah dibangun?", "Live functional prototype: Mobile App Android/Web (bjb & Kuningan), Core API, AI Engine, Dashboard.", "LOLOS"),
        ("5. Technical", "Dapatkah dijelaskan bagaimana sistem menghasilkan output?", "Input JSON -> Rule Engine (13 rules) + Random Forest (29 fitur) + GraphSAGE lookup -> Risk Score.", "LOLOS"),
        ("6. Testing", "Apakah ada hasil pengujian terhadap fungsi & asumsi utama?", "Unit test test_rule_engine.py (5/5 PASS), API integration test, evaluasi test set (64.122 sampel).", "LOLOS"),
        ("7. Impact", "Apakah angka dampak memiliki sumber dan metode perhitungan?", "Waktu LTKM dari 3 hari ke 3 detik (otomasi template); penyelamatan dana 100% pada transaksi kritis.", "LOLOS"),
        ("8. Market", "Apakah sudah ada bukti validasi dari calon offtaker nyata?", "Uji keselarasan solusi bersama Bank bjb (TC-BJB-01 s/d 04) dan Bank Kuningan (TC-KNG-01 s/d 04).", "LOLOS"),
        ("9. Differentiation", "Apakah tim memahami alternatif eksisting dan keunikan solusinya?", "Benchmark terhadap FDS Rule statis, Excel manual, dan solusi enterprise global (SAS/Actimize).", "LOLOS"),
        ("10. Team", "Apakah peran, kapabilitas, dan ownership aktual anggota tim jelas?", "Kepemilikan kode jelas: Rifki (AI/Lead), Billy (Security/CBS), Aam (Frontend/UI), Desta (Backend/LTKM).", "LOLOS"),
        ("11. Continuation", "Apakah tim tahu milestone berikutnya dan apa yang dibutuhkan?", "Proposal pilot 3 bulan: Bulan 1 (Sandbox), Bulan 2 (Shadow deployment Mode B), Bulan 3 (Audit SKAI).", "LOLOS"),
        ("12. Transparency", "Apakah deck membedakan secara tegas hal yang sudah jadi vs rencana?", "Status Live Prototype diberi label transparan; batasan runtime GNN dan integrasi live CBS diungkapkan.", "LOLOS")
    ]
    t_el = doc.add_table(rows=len(eligibility_data)+1, cols=4)
    t_el.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_el)
    el_headers = ["Aspek Uji", "Pertanyaan Kelayakan Resmi PIDI", "Bukti Nyata Solusi Crypto-Sentinel", "Status"]
    el_w = [Inches(1.3), Inches(2.2), Inches(2.8), Inches(0.7)]
    for j, h in enumerate(el_headers):
        cell = t_el.cell(0, j)
        cell.width = el_w[j]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(eligibility_data):
        for j, val in enumerate(row):
            cell = t_el.cell(i+1, j)
            cell.width = el_w[j]
            bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if j == 0:
                r.bold = True
            elif j == 3:
                r.bold = True
                r.font.color.rgb = RGBColor(22, 163, 74)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 4. RINGKASAN EKSEKUTIF SLIDE DECK (12 SLIDES)
    add_heading_styled(doc, "Ringkasan Eksekutif Slide Deck (12 Slides Resmi PIDI Digdaya)", level=1)
    deck_summary = [
        ("1", "Solution at a Glance", "Identitas, BPD/BPR, Single-sentence Value Prop", "Profil BPD/BPR, Status Live Functional Prototype"),
        ("2", "Problem & Why It Matters", "Kerugian Fraud Transfer, Sindikat Mule, Kripto", "OJK IASC Rp 9,1T, Kasus BI-FAST BPD ratusan miliar"),
        ("3", "Validation & Root Cause", "FDS Lama Buta Topologi Graf & Banjir False Alarm", "Validasi Bank bjb (TC-BJB-01) & Bank Kuningan (TC-KNG-01)"),
        ("4", "Solution & Core Use Case", "Alur Intersepsi Pre-Auth & Investigasi Post-Auth", "Mobile Transfer -> Breaker (<25ms) -> CMS -> goAML LTKM"),
        ("5", "Value Proposition", "Komparasi vs Rule-Only & Enterprise Global", "Biaya hemat 70%, XAI SHAP, Kepatuhan POJK 8/2023 & UU PDP"),
        ("6", "Prototype State", "Kesiapan Sistem 4-Layer, Fungsi Nyata, Batasan", "Flutter Mobile, Core API FastAPI, AI Engine, React RBAC"),
        ("7", "How Technology Works", "Arsitektur 3-Layer, Input-Process-Output AI", "max(0.6*GNN + 0.4*Rule, Rule_Floor), GNNExplainer MI"),
        ("8", "Technical Testing", "Benchmark Model, Latensi, Uji Keamanan", "Akurasi 99.98%, AUC 0.9997, FPR 0.002%, Latensi 5.67ms"),
        ("9", "Impact & Effectiveness", "Penyelamatan Dana, Reduksi FP, Efisiensi STR", "Waktu LTKM dari 3 hari ke 3 detik, Zero-loss Circuit Breaker"),
        ("10", "Market Validation", "Hasil Uji Lapangan Bank bjb & Bank Kuningan", "Notulensi 8 Test Cases (Mitigasi PII Masking, FP & Audit Log)"),
        ("11", "Adoption Path", "Roadmap Pilot 3 Bulan, Biaya, Mode A & B", "Mode A (Pre-Auth) & Mode B (Read-Only CBS), TCO Terjangkau"),
        ("12", "Team Readiness", "Kepemilikan Peran Tim EXPRESSO S1251", "Ownership AI/Lead, Security/CBS, Frontend/UI, Backend/LTKM")
    ]
    t_sum = doc.add_table(rows=len(deck_summary)+1, cols=4)
    t_sum.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_sum)
    s_headers = ["#", "Judul Slide Resmi", "Fokus Konten Utama", "Bukti & Evidence Kunci"]
    s_w = [Inches(0.4), Inches(2.2), Inches(2.3), Inches(2.1)]
    for j, h in enumerate(s_headers):
        cell = t_sum.cell(0, j)
        cell.width = s_w[j]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(deck_summary):
        for j, val in enumerate(row):
            cell = t_sum.cell(i+1, j)
            cell.width = s_w[j]
            bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8)
            if j <= 1:
                r.bold = True

    doc.add_page_break()

    # 5. KONTEN DETAIL 12 SLIDE RESMI
    add_heading_styled(doc, "Materi Lengkap 12 Slide Presentasi & Naskah Presenter", level=1)
    
    slides_data = [
        {
            "num": 1,
            "title": "SLIDE 1: SOLUTION AT A GLANCE",
            "ref": "Progress Report: Bab 1.1 Profil Proyek  |  Blueprint: Section 00 Overview",
            "visual": "Enterprise Banking Dark Navy (#09132E), Badge: FIELD-VALIDATED PROTOTYPE / CONTROLLED SANDBOX CANDIDATE, Mockup Device Ganda (Mobile Banking Android + Forensic Dashboard di Laptop).",
            "headline": "Crypto-Sentinel 2026",
            "subheadline": "Next-Generation Hybrid FDS & AML Middleware for Resilient Banking Sovereignty.",
            "points": [
                ("Problem Statement", "Bank Pembangunan Daerah (BPD) dan BPR menjadi sasaran empuk sindikat pencucian uang modern (mule networks, smurfing, dan crypto illicit outflow) karena keterbatasan sistem deteksi konvensional."),
                ("Target User & Offtaker", "Offtaker: PT BPR Kuningan (Perseroda) & PT Bank Pembangunan Daerah Jawa Barat dan Banten Tbk (Bank bjb). Pengguna: Analis AML, Compliance Officer (MLRO), Pengawas OJK/BI."),
                ("Golden Positioning Statement", "Field-validated, plug-and-play FDS security middleware prototype, ready for controlled sandbox deployment and pilot hardening — not yet production-integrated with a bank CBS."),
                ("Single-Sentence Value Proposition", "Crypto-Sentinel 2026 adalah middleware keamanan FDS berbasis AI hibrida yang bertindak sebagai Smart Circuit Breaker, menghentikan transaksi sindikat pencucian uang dalam hitungan milidetik sebelum saldo keluar dari bank, dilengkapi otomasi draf pelaporan resmi PPATK.")
            ],
            "script": "Selamat pagi Dewan Juri yang terhormat. Saya Rifki Firmansyah, mewakili Tim EXPRESSO S1251. Kami mempersembahkan Crypto-Sentinel 2026. Kami membangun field-validated FDS security middleware berbasis AI hibrida yang dirancang khusus untuk melindungi BPR dan Bank Daerah dari ancaman pembobolan transfer cepat dan sindikat rekening mule. Sistem kami adalah live prototype yang telah diuji langsung bersama bank mitra, mengintegrasikan mobile banking nasabah, core banking simulator, engine graf relasional, hingga otomasi draf pelaporan resmi ke PPATK."
        },
        {
            "num": 2,
            "title": "SLIDE 2: PROBLEM & WHY IT MATTERS",
            "ref": "Progress Report: Bab 1.1 Latar Belakang & Kasus BPD  |  Blueprint: Hero Stats & Scope",
            "visual": "High Urgency Security Alert, 3 Stat Callouts Raksasa di bagian atas, Diagram Alur Serangan BI-FAST di bagian bawah.",
            "headline": "Ancaman Nyata: Eksploitasi Transfer Cepat & Pencucian Uang Lintas Kanal",
            "subheadline": "Mengapa masalah ini mendesak dan menimbulkan kerugian sistemik pada bank-bank daerah.",
            "points": [
                ("Rp 9,1 Triliun", "Akumulasi kerugian transaksi keuangan ilegal nasional yang dihimpun Satgas PASTI / OJK IASC."),
                ("Rp 800+ Miliar", "Kerugian riil insiden fraud transfer antarbank cepat (kasus BI-FAST) yang dialami bank-bank daerah di Indonesia akibat eksploitasi celah middleware."),
                ("> 85% Rekening Penampung (Mule)", "Ditempatkan di bank menengah dan BPR lokal dengan sistem pengawasan pasca-transaksi yang lambat, sebelum dilarikan ke bursa aset kripto dalam hitungan menit."),
                ("Konsekuensi Nyata Bank", "Pemberitahuan terlambat (T+1/T+2 dana sudah hilang); Risiko sanksi kepatuhan POJK No. 8/2023 (Penerapan Strategi Anti-Fraud); Kehilangan reputasi dan beban ganti rugi nasabah.")
            ],
            "script": "Mengapa masalah ini sangat mendesak? OJK dan PPATK mencatat kerugian kejahatan finansial telah menembus angka triliunan rupiah. Salah satu kasus paling fatal menimpa BPD dengan kebocoran dana transfer ratusan miliar rupiah. Sindikat kejahatan masa kini mengeksploitasi kecepatan transfer digital; mereka memecah dana hasil kejahatan melalui skema smurfing ke puluhan rekening mule di bank daerah, lalu mengurasnya ke bursa kripto dalam hitungan menit. Ketika nasabah baru menyadari saldonya hilang di pagi hari, dana tersebut sudah lenyap di blockchain tanpa bisa ditarik kembali."
        },
        {
            "num": 3,
            "title": "SLIDE 3: PROBLEM VALIDATION & ROOT CAUSE",
            "ref": "Progress Report: Bab 5.1 Temuan Validasi Stakeholder  |  Blueprint: Section 01 Celah Deteksi",
            "visual": "Root Cause Analysis, Kolom Kiri Hasil Validasi Wawancara Bank, Kolom Kanan Diagram 3 Akar Masalah Utama.",
            "headline": "Validasi Stakeholder: Mengapa Sistem FDS Konvensional Gagal?",
            "subheadline": "Temuan langsung dari uji keselarasan solusi bersama tim IT & Kepatuhan Bank bjb dan Bank Kuningan.",
            "points": [
                ("Akar Masalah #1: Buta Topologi Graf (Relational Blindness)", "FDS konvensional hanya mengecek transaksi tunggal. Jika transfer Rp 4,9 juta (di bawah limit Rp 5 juta), sistem menganggap normal. Sistem buta bahwa ada 10 akun serentak mentransfer ke satu rekening penampung yang sama (fan-in mule ring)."),
                ("Akar Masalah #2: Banjir Peringatan Palsu (High False Positives)", "Aturan statis memicu alarm pada transaksi sah (pencairan Bansos massal atau pembayaran SPP sekolah rutin), menyebabkan alert fatigue bagi analis."),
                ("Akar Masalah #3: Deteksi Pasif Pasca-Transaksi (Post-Factum Lag)", "FDS eksisting bekerja secara batch/post-audit (T+1). Uang sudah keluar dari bank sebelum investigasi dimulai."),
                ("Kesimpulan Validasi", "Aturan Statis Tunggal + Tanpa Pemetaan Graf Relasi = Kerugian Dana Tidak Terbendung.")
            ],
            "script": "Ketika kami memvalidasi masalah ini bersama praktisi IT dan Kepatuhan Bank bjb serta Bank Kuningan, kami menemukan akar masalahnya: FDS yang ada saat ini bekerja dengan kacamata kuda. Sistem lama hanya mengecek parameter statis satu per satu. Sindikat kriminal sangat paham celah ini; mereka memecah transaksi bernominal kecil agar lolos dari aturan. FDS konvensional buta terhadap graf relasi rekening penampung, dan di sisi lain justru menghasilkan ribuan alarm palsu pada transaksi sah seperti Bansos. Akibatnya, analis kewalahan dan kejahatan yang sesungguhnya tetap lolos."
        },
        {
            "num": 4,
            "title": "SLIDE 4: SOLUTION & CORE USE CASE",
            "ref": "Progress Report: Bab 1.2 Dual-Mode Engine  |  Blueprint: Section 01 Flow Diagram Intersepsi",
            "visual": "Core Use Case Process Flow, 4 Kartu Horizontal Terhubung dari Sisi Nasabah hingga Regulator.",
            "headline": "Cara Kerja Solusi: Intersepsi Sub-Detik hingga Pelaporan Resmi",
            "subheadline": "Core use case penanganan transaksi mencurigakan secara instan dan terstruktur.",
            "points": [
                ("Step 1: User Inisiasi Transfer", "Nasabah/pelaku memicu transfer dana via Mobile Banking berotentikasi SNAP BI (HMAC-SHA256)."),
                ("Step 2: Evaluasi Hibrida Sub-Detik (<25 ms)", "Request dicegat di middleware sebelum mutasi saldo dieksekusi di Core Banking. AI Engine menghitung skor risiko gabungan (Perilaku, Tabular ML, dan Topologi Relasional GNN)."),
                ("Step 3: Smart Circuit Breaker Seketika", "Jika Skor >= 85, transaksi langsung di-BLOCK otomatis, saldo nasabah aman 100%, dan notifikasi dikirim ke nasabah secara profesional."),
                ("Step 4: Investigasi & Otomasi goAML LTKM", "Analis membuka kasus di Forensic Dashboard, memverifikasi relasi sindikat via GNNExplainer 3-hop, dan mencetak draf resmi LTKM PPATK format goAML dalam 3 detik.")
            ],
            "script": "Bagaimana Crypto-Sentinel menyelesaikan masalah ini? Mari kita lihat use case intinya: Saat transaksi transfer dipicu di mobile banking, request dicegat di layer middleware sebelum menyentuh core banking. Dalam waktu kurang dari 25 milidetik, AI engine kami menghitung skor risiko hibrida. Jika terdeteksi indikasi berbahaya seperti pengurasan saldo ke bursa kripto, transaksi langsung di-BLOCK seketika—saldo nasabah tidak berkurang sama sekali. Detik itu juga, alert masuk ke dashboard analis, subgraf sindikatnya divisualisasikan, dan draf laporan resmi goAML PPATK siap diterbitkan."
        },
        {
            "num": 5,
            "title": "SLIDE 5: VALUE PROPOSITION & DIFFERENTIATION",
            "ref": "Progress Report: Bab 1.3 Keunggulan Kompetitif  |  Blueprint: Section 00 & 03 Uniqueness",
            "visual": "Matriks Komparasi Multi-Dimensi terhadap Proses Manual, Rule Internal, Vendor Global SAS/Actimize.",
            "headline": "Diferensiasi & Keunggulan terhadap Alternatif Eksisting",
            "subheadline": "Mengapa Crypto-Sentinel adalah pilihan paling rasional, aman, dan terjangkau bagi bank daerah.",
            "points": [
                ("Kecepatan Intersepsi", "Manual: T+2  |  Rule Internal: Jam/Menit  |  Enterprise Global: <50ms  |  Crypto-Sentinel: Sub-detik (<25 ms Pre-Auth Breaker)."),
                ("Deteksi Sindikat Mule", "Manual: Buta Total  |  Rule Internal: Gagal  |  Enterprise Global: Ada (perlu server GPU raksasa)  |  Crypto-Sentinel: Unggul (GraphSAGE GNN 562K Nodes)."),
                ("Transparansi Keputusan (XAI)", "Manual: Narasi manual  |  Rule Internal: Pesan rule kaku  |  Enterprise Global: Black Box  |  Crypto-Sentinel: Tinggi (SHAP Factors + Subgraf GNNExplainer)."),
                ("Kepatuhan Regulasi Lokal", "Manual: Rentan sanksi  |  Rule Internal: Parsial  |  Enterprise Global: Mahal (>Rp 500 Juta add-on)  |  Crypto-Sentinel: Native (POJK 8/2023, LTKM goAML, APOLO)."),
                ("Privasi Data (UU PDP)", "Manual: Terbuka penuh  |  Rule Internal: Tanpa masking  |  Enterprise Global: Konfigurasi rumit  |  Crypto-Sentinel: Native Privacy Masking (UU PDP No. 27/2022)."),
                ("Total Cost of Ownership", "Enterprise Global: Miliaran Rupiah/tahun  |  Crypto-Sentinel: Efisiensi Tinggi (Hemat hingga 70% TCO).")
            ],
            "script": "Di pasar saat ini, bank daerah terjepit di antara dua pilihan yang tidak ideal: menggunakan FDS aturan internal atau spreadsheet manual yang terbukti gagal membendung pencucian uang modern, atau membeli software enterprise global bernilai miliaran rupiah yang mahal dan tidak memahami regulasi lokal Indonesia. Crypto-Sentinel hadir memberikan diferensiasi nyata: kami menghadirkan kemampuan deteksi jaringan setara software kelas dunia melalui GraphSAGE GNN, namun dengan arsitektur komputasi yang ringan tanpa server GPU mahal di runtime, dan sudah terpasang modul kepatuhan lokal seperti standar pelaporan goAML PPATK dan UU PDP."
        },
        {
            "num": 6,
            "title": "SLIDE 6: PROTOTYPE & CURRENT PRODUCT STATE",
            "ref": "Progress Report: Bab 2.1 Status Komponen Terpasang  |  Blueprint: Section 02 Tech Stack & 06 DB",
            "visual": "Live Product Evidence Grid & 3-Layer Architecture Diagram, 4 Tangkapan Layar Aplikasi Nyata.",
            "headline": "Status Kesiapan Produk: Field-Validated Prototype",
            "subheadline": "Arsitektur 3-Layer FDS Console & 8 Fungsi Inti Terpenuhi.",
            "points": [
                ("Golden Positioning", "Field-validated, plug-and-play FDS security middleware prototype, ready for controlled sandbox deployment and pilot hardening — not yet production-integrated with a bank CBS."),
                ("Layer 1: FDS Control Tower", "Telemetri operasional, throughput transaksi, dana tertahan (Blocked Value), latensi p50/p95, dan status data freshness."),
                ("Layer 2: Alert & Case Management", "Pusat investigasi, rekening tersamar (UU PDP), dekomposisi skor risiko (Rule/ML/GNN), SHAP reasons, subgraf GNNExplainer, maker-checker, dan draft LTKM goAML."),
                ("Layer 3: Intelligence & Governance", "Pengaturan risk appetite, kalibrasi threshold, deteksi drift & FP, manajemen multi-tenant, dan audit trail persisten."),
                ("Source-of-Truth Badges", "LIVE · Sentinel API, LIVE · Core Banking API, SYNTHETIC · PaySim, DEMO FIXTURE, dan ERROR (menghilangkan silent mock fallback)."),
                ("Transparansi Batasan Saat Ini", "Validasi dilakukan pada Core Banking Simulator terstandarisasi SNAP BI; integrasi ke CBS nyata dijadwalkan pada Fase Pilot dengan penandatanganan NDA, DPA, dan audit independen.")
            ],
            "script": "Sesuai prinsip 'Built over Planned', Crypto-Sentinel adalah Field-Validated Prototype yang memenuhi 8 fungsi inti FDS perbankan. Konsol dashboard kami dibangun di atas arsitektur 3-Layer: Control Tower untuk telemetri sistem, Case Management untuk investigasi subgraf dan penerbitan draf LTKM PPATK, serta Intelligence Layer untuk kalibrasi regulasi POJK. Sistem kami menerapkan tata kelola data yang transparan dengan badge sumber data aktif, siap memasuki fase uji coba sandbox terkontrol bersama bank mitra."
        },
        {
            "num": 7,
            "title": "SLIDE 7: HOW THE TECHNOLOGY WORKS (AI & SYSTEM LOGIC)",
            "ref": "Progress Report: Bab 3.1 Spesifikasi AI Hibrida  |  Blueprint: Section 03 Indikator & 05 AI Dev",
            "visual": "Deep-Tech Architecture Diagram, Input-Process-Output AI, Formula Fusi Hibrida, XAI Subgraf.",
            "headline": "Arsitektur AI: Spesifikasi Model & Logika Pemrosesan Sistem",
            "subheadline": "Bukan sekadar 'Powered by AI'—kombinasi deterministik dan pemetaan graf relasional.",
            "points": [
                ("Input Model", "Payload transaksi JSON (nominal, saldo awal/akhir pengirim & penerima, timestamp, IP geolocation, Device ID, kode tujuan ISO 20022)."),
                ("Sumber Data", "Dataset PaySim diaugmentasi menjadi 320.606 transaksi + 12.393 edge cases lokal Indonesia (Bansos, SPP, QRIS, Kripto outflow)."),
                ("Pemrosesan Data", "Ekstraksi 29 fitur tabular + graf transaksi NetworkX (562K nodes) + offline GraphSAGE GNN (PyG) embeddings di Google Colab GPU T4."),
                ("Output Model", "Skor Risiko Kontinu (0-100), Klasifikasi Keputusan (ALLOW, REVIEW, BLOCK), dan kontribusi fitur SHAP value."),
                ("Formula Fusi Hibrida", "hybrid_score = 0.6 * gnn_score + 0.4 * rule_score  ;  final_score = max(hybrid_score, rule_score). Rule Engine bertindak sebagai floor safety signal."),
                ("Keterbatasan & Oversight", "GraphSAGE offline lookup; jika entitas baru, Rule Engine mengunci fraud secara deterministik. Human-in-the-loop: skor 60-84 wajib verifikasi manual analis.")
            ],
            "script": "Bagaimana teknologinya bekerja? Kami tidak sekadar menempelkan stempel 'Powered by AI'. Model kami dilatih menggunakan 320 ribu data transaksi yang telah diperkaya 12 ribu kasus perbankan Indonesia. Sinyal keputusan digabungkan melalui formula hibrida: Rule Engine 13 indikator menyaring anomali teknis dan bertindak sebagai safety floor signal; Random Forest memproses fitur nominal; dan GraphSAGE GNN memetakan relasi jaringan rekening penampung. Yang terpenting: sistem kami menerapkan prinsip Human-in-the-Loop. Transaksi berskor menengah tidak diblokir sepihak, melainkan dialihkan ke antrean verifikasi analis kepatuhan."
        },
        {
            "num": 8,
            "title": "SLIDE 8: TECHNICAL TESTING & PERFORMANCE",
            "ref": "Progress Report: Bab 4.1 Benchmark Model AI  |  Blueprint: Hero Stats (<50ms, >95% F1)",
            "visual": "Testing Matrix, Metric Cards (Akurasi, AUC, FPR, Latensi), Kurva Evaluasi ROC-AUC.",
            "headline": "Pengujian Teknis & Tolok Ukur Kinerja Sistem",
            "subheadline": "Validasi performa model machine learning, latensi komputasi, dan pengujian integrasi.",
            "points": [
                ("Kinerja AI / ML (Test Set: 61.643 Sampel)", "Akurasi: 99.98%  |  Precision: 99.94%  |  Recall Fraud: 99.88%  |  F1-Score: 99.91%  |  ROC-AUC: 0.9997."),
                ("Mitigasi False Positive Rate (FPR)", "0.002% (Hanya 1 false alarm dari 60.000 transaksi legitimate perbankan sah)."),
                ("Pencegahan Missed Fraud (FNR)", "0.122% (Hanya 2 fraud terlewat dari 1.643 kasus fraud pada test set)."),
                ("Inference Response Time Platform", "Rata-rata 5.67 ms (p95: 9.05 ms · p99: 12.23 ms pada CPU lokal sandbox)."),
                ("Otomasi Dokumen Regulasi", "Waktu terbit dokumen draf goAML LTKM PPATK: 3 Detik (Eliminasi 95% langkah administrasi manual)."),
                ("Keamanan & Integritas", "Verifikasi tanda tangan SNAP BI HMAC-SHA256; Audit trail persisten pada tabel AuditLog dengan proteksi RBAC."),
                ("Unit & Benchmark Test Suite", "test_rule_engine.py: 5/5 PASS  |  eval_official_benchmark.py: PASS (308.213 PaySim terverifikasi).")
            ],
            "script": "Kinerja teknis Crypto-Sentinel telah teruji secara kuantitatif melalui benchmark resmi pada 308 ribu transaksi PaySim. Model hibrida kami mencapai akurasi 99.98%, presisi 99.94%, recall fraud 99.88%, dan skor ROC-AUC 0.9997. Yang paling krusial bagi bank: False Positive Rate kami ditekan hingga 0.002%—artinya hanya 1 dari 60 ribu transaksi sah yang terganggu. Dari sisi performa, latensi inferensi rata-rata kami hanya 5.67 milidetik dengan p95 di bawah 10 milidetik, jauh melampaui standar kecepatan industri perbankan."
        },
        {
            "num": 9,
            "title": "SLIDE 9: IMPACT & EVIDENCE OF EFFECTIVENESS",
            "ref": "Progress Report: Bab 1.3 Transformasi Operasional  |  Blueprint: Section 04 Otomasi LTKM PPATK",
            "visual": "Tabel Baku Pengukuran Dampak Resmi PIDI: KPI -> Baseline -> Hasil Saat Ini -> Target -> Evidence.",
            "headline": "Dampak Nyata & Efektivitas Operasional",
            "subheadline": "Pengukuran dampak terstruktur antara kondisi eksisting (baseline) dan hasil implementasi.",
            "points": [
                ("Penyelamatan Dana Fraud", "Baseline: 0% (Dana keluar dahulu, baru diselidiki di T+1) -> Hasil Saat Ini: 100% Tertahan pada simulasi smurfing kritis -> Target: >=90% Dana Terselamatkan."),
                ("Waktu Terbit Draf LTKM PPATK", "Baseline: 2 s.d. 3 Hari Kerja manual -> Hasil Saat Ini: 3 Detik sekali klik -> Target: < 1 Menit terkirim."),
                ("False Positive Interruption", "Baseline: Tinggi (nasabah komplain saat gajian) -> Hasil Saat Ini: 0.002% FPR (1 / 60.000) -> Target: < 0.5% FPR."),
                ("Inference Response Time", "Baseline: Lambat (>500ms atau batch post-audit) -> Hasil Saat Ini: 5.67 ms (p95: 9.05ms) -> Target: < 25 ms."),
                ("Waktu Triage Investigasi Analis", "Baseline: Berjam-jam menelusuri mutasi rekening koran -> Hasil Saat Ini: < 2 Menit via GNNExplainer 3-hop -> Target: < 5 Menit per kasus."),
                ("Kesiapan Audit Regulasi", "Baseline: Log tersebar rawan manipulasi -> Hasil Saat Ini: 100% Tercatat di tabel immutable audit_logs -> Target: Audit-Ready OJK & PPATK.")
            ],
            "script": "Dampak apa yang dihasilkan oleh Crypto-Sentinel? Kami mengukurnya melalui metodologi ketat: Pertama, Penyelamatan Dana. Berkat Smart Circuit Breaker, potensi kerugian fraud ditekan hingga 100% pada skenario kritis karena dana dicegat sebelum keluar dari bank. Kedua, Efisiensi Kepatuhan. Penyusunan draf Laporan Transaksi Keuangan Mencurigakan (LTKM) ke PPATK yang biasanya memakan waktu 2 hingga 3 hari kerja manual, kini terpangkas menjadi hanya 3 detik. Ketiga, Pengurangan False Positive. Transaksi sah nasabah tetap lancar dengan rasio kesalahan di bawah 0.01%, menjaga kepuasan nasabah sekaligus kepatuhan audit regulasi."
        },
        {
            "num": 10,
            "title": "SLIDE 10: MARKET / USER / OFFTAKER VALIDATION",
            "ref": "Progress Report: Bab 5.2 Notulensi Uji 8 Kasus  |  Blueprint: Section 07 Roadmap Kemitraan",
            "visual": "Offtaker Evidence Cards, Notulensi Uji Coba Bank bjb & Bank Kuningan, 3 Pertanyaan Kunci Pasar.",
            "headline": "Validasi Offtaker & Pembelajaran Lapangan",
            "subheadline": "Interaksi nyata, pengujian bersama calon pengguna, dan perbaikan berbasis masukan bank.",
            "points": [
                ("Who Needs It? (Kebutuhan Nyata)", "Divisi Kepatuhan (Compliance/APU-PPT) dan Divisi TI di 26 Bank Pembangunan Daerah (BPD) dan >1.400 BPR di Indonesia."),
                ("What Have You Learned? (Uji Coba 25 Agustus 2026)", "Bank bjb & Bank Kuningan: Catatan anonimisasi data PII direspons dengan penambahan Privacy Masking; Risiko false positive direspons dengan pemisahan antrean REVIEW manual (skor 60-84); Kebutuhan tata kelola direspons dengan penambahan tabel AuditLog dan CMS."),
                ("What Signal of Adoption Exists?", "Bank Kuningan: Minat pilot terkonfirmasi untuk pengujian transaksi BPR di sandbox terkontrol; Bank bjb: Validasi keselarasan format laporan LTKM goAML PPATK dinyatakan PASS (TC-BJB-03).")
            ],
            "script": "Kami memegang teguh prinsip 'Validation over Assumption'. Pada tanggal 25 Agustus lalu, kami menguji langsung sistem ini bersama praktisi IT dan Kepatuhan Bank bjb serta Bank Kuningan. Kami belajar banyak hal berharga: bank meminta perlindungan data pribadi nasabah dan menolak pemblokiran otomatis pada transaksi yang meragukan. Masukan tersebut langsung kami eksekusi: kami membangun Privacy Masking sesuai UU PDP, memisahkan antrean review manual untuk mitigasi false positive, dan mengintegrasikan tabel audit log permanen. Kedua mitra mengonfirmasi bahwa format pelaporan dan mekanisme deteksi kami sangat relevan dengan kebutuhan operasional mereka."
        },
        {
            "num": 11,
            "title": "SLIDE 11: ADOPTION & SUSTAINABILITY PATH",
            "ref": "Progress Report: Bab 6.1 Roadmap Pilot 3 Bulan  |  Bank Integration Kit: docs/BANK_INTEGRATION_KIT.md  |  Blueprint: Section 07 Roadmap Pelaksanaan",
            "visual": "Roadmap Eksekusi 3 Bulan, Diagram 2 Mode Integrasi Fleksibel (Mode A Pre-Auth & Mode B Read-Only).",
            "headline": "Jalur Adopsi Praktis & Model Keberlanjutan",
            "subheadline": "Arsitektur Plug-and-Play by Integration Pattern & Roadmap Pilot 3 Bulan.",
            "points": [
                ("Konsep Plug-and-Play by Integration Pattern", "Tidak mengubah core banking lama; Bank menghubungkan adapter transaksi (REST/Webhook/CDC) ke canonical FDS API."),
                ("Bank Integration Kit (docs/BANK_INTEGRATION_KIT.md)", "Environment schema .env.local (APP_MODE=live|demo|hybrid), canonical schema mapping, dan mTLS/HMAC config."),
                ("Mode A: Pre-Authorization Gateway", "Terpasang di API Gateway Mobile Banking / SNAP BI untuk memutus transaksi penipuan secara real-time."),
                ("Mode B: Post-Transaction Monitoring", "Menerima salinan event transaksi read-only dari CDC database BPR Kuningan tanpa mengubah atau membebani Core Banking lama."),
                ("Milestone Pilot Sandbox (Bulan 1-2)", "Adapter contract per-bank, HMAC/request auth, backend RBAC & tenant isolation, Privacy Masking, audit log persisten, fail-safe (allow/hold/deny)."),
                ("Milestone Menuju Produksi (Bulan 3+)", "Security pentest independen, load test p95/p99, high availability/DR, formal DPA approval, dan change approval bank."),
                ("Model Keberlanjutan Finansial", "Skema berlangganan SaaS/On-Premise bertingkat (Tier BPR Rp 5 Juta/bln, Tier BPD Rp 25 Juta/bln), menghemat hingga 70% TCO dibanding lisensi FDS global.")
            ],
            "script": "Bagaimana solusi ini akan diadopsi secara nyata? Kami menerapkan prinsip 'Plug-and-play by integration pattern'—artinya bank tidak perlu merombak core banking lama. Melalui Bank Integration Kit yang kami siapkan, bank dapat memilih Mode A untuk perlindungan seketika di kanal digital atau Mode B untuk pemantauan pasca-transaksi secara read-only. Kami membagi roadmap eksekusi secara disiplin: pengujian sandbox dan audit privasi di bulan pertama, shadow deployment di bulan kedua, dan audit SKAI di bulan ketiga, didukung skema biaya yang sangat terjangkau bagi BPR dan BPD."
        },
        {
            "num": 12,
            "title": "SLIDE 12: TEAM & EXECUTION READINESS",
            "ref": "Progress Report: Bab 1.1 Tim Pengembang  |  Blueprint: Section 08 Tim EXPRESSO UNIKU",
            "visual": "4 Profil Anggota Tim EXPRESSO S1251 dengan Kepemilikan Pekerjaan Aktual (Role + Capability + Actual Ownership).",
            "headline": "Tim EXPRESSO (S1251) & Kesiapan Eksekusi",
            "subheadline": "Kolaborasi lintas disiplin AI, keamanan siber perbankan, dan rekayasa perangkat lunak enterprise.",
            "points": [
                ("Rifki Firmansyah (Team Lead & AI Architect)", "Ownership: Arsitektur model hibrida, training GraphSAGE GNN di Colab GPU T4, Rule Engine 13 indikator, dan memimpin uji coba bank mitra."),
                ("Billy Jonathan (Cyber Security & Core Banking)", "Ownership: Keamanan SNAP BI HMAC-SHA256, Core Banking Simulator Gateway (Expresso API), dan validasi anti-tampering."),
                ("Aam Setiana (Frontend Engineer & Product Analyst)", "Ownership: Forensic Compliance Dashboard (React 18 + Vite), visualisasi subgraf GNNExplainer, dan RBAC 3-tier."),
                ("Desta Erlangga (Backend & Integration Engineer)", "Ownership: Endpoint REST API, generator draf dokumen resmi LTKM PPATK (goAML), dan tabel AuditLog persisten."),
                ("Mentor Industri", "Bayu Ferdian, MBA., CIP. (CEO Gizalab — Praktisi Customer Experience & Product Strategy)."),
                ("Closing Hook", "Crypto-Sentinel 2026: Siap Melangkah dari Prototype Menuju Pilot Nyata untuk Menjaga Kedaulatan Finansial Perbankan Indonesia.")
            ],
            "script": "Di balik inovasi ini, ada komitmen eksekusi penuh dari Tim EXPRESSO S1251. Kami membagi kepemilikan teknis secara nyata: saya sendiri mengawal arsitektur AI dan strategi produk; Billy Jonathan mengamankan arsitektur perbankan dan SNAP BI; Aam Setiana merancang pengalaman visual investigasi forensik di dashboard; dan Desta Erlangga memastikan integrasi backend dan kepatuhan PPATK berjalan andal. Didukung arahan mentor industri serta validasi nyata dari Bank bjb dan Bank Kuningan, kami membuktikan bahwa tim kami siap mengeksekusi solusi ini ke tahap berikutnya. Terima kasih, kami siap untuk sesi tanya jawab."
        }
    ]

    for s in slides_data:
        add_heading_styled(doc, s["title"], level=2)
        
        # Reference Tag Box
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_before = Pt(1)
        p_ref.paragraph_format.space_after = Pt(4)
        r_ref_tag = p_ref.add_run("📌 Rujukan Teknis: ")
        r_ref_tag.bold = True
        r_ref_tag.font.size = Pt(8.5)
        r_ref_tag.font.color.rgb = RGBColor(2, 132, 199)
        r_ref_val = p_ref.add_run(s["ref"])
        r_ref_val.font.size = Pt(8.5)
        r_ref_val.font.color.rgb = RGBColor(100, 116, 139)

        # Blueprint Visual Guide
        p_vis = doc.add_paragraph()
        p_vis.paragraph_format.space_before = Pt(1)
        p_vis.paragraph_format.space_after = Pt(4)
        r_vis_tag = p_vis.add_run("Panduan Visual Blueprint: ")
        r_vis_tag.bold = True
        r_vis_tag.font.size = Pt(8.5)
        r_vis_val = p_vis.add_run(s["visual"])
        r_vis_val.font.size = Pt(8.5)
        r_vis_val.font.italic = True

        # Headline & Subheadline
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(2)
        p_h.paragraph_format.space_after = Pt(1)
        r_h = p_h.add_run(s["headline"])
        r_h.bold = True
        r_h.font.size = Pt(11)
        r_h.font.color.rgb = RGBColor(15, 23, 42)
        
        p_sub = doc.add_paragraph()
        p_sub.paragraph_format.space_before = Pt(0)
        p_sub.paragraph_format.space_after = Pt(4)
        r_sub = p_sub.add_run(s["subheadline"])
        r_sub.font.italic = True
        r_sub.font.size = Pt(9.5)
        r_sub.font.color.rgb = RGBColor(100, 116, 139)

        # Points
        for k, v in s["points"]:
            p_pt = doc.add_paragraph(style='List Bullet')
            p_pt.paragraph_format.space_before = Pt(1)
            p_pt.paragraph_format.space_after = Pt(2)
            r_k = p_pt.add_run(k + ": ")
            r_k.bold = True
            r_k.font.size = Pt(9.5)
            r_v = p_pt.add_run(v)
            r_v.font.size = Pt(9.5)

        # Speaker Script Callout Box
        add_callout(doc, "🎙️ Naskah Presenter (Speaker Script — Waktu: ~25 Detik):", s["script"], bg="F0F9FF", border_color="0284C7")
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    doc.add_page_break()

    # 6. PANDUAN TANYA JAWAB (Q&A DEFENSE SHEET)
    add_heading_styled(doc, "Panduan Menjawab Pertanyaan Kritis Juri (Q&A Defense Sheet)", level=1)
    qa_list = [
        ("Q1: Apakah sistem Anda melakukan full message-passing GNN di setiap transaksi? Bagaimana latensinya bisa <25ms?",
         "Tidak. Menjalankan deep message-passing online pada setiap transaksi perbankan akan menimbulkan latensi tinggi dan beban GPU yang mahal. Kami mengadopsi pola enterprise: GraphSAGE dilatih offline untuk mengekstrak embeddings 32-dimensi. Di runtime API, kami melakukan fast embedding lookup yang digabungkan dengan Random Forest dan Rule Engine pada CPU ringan. Inilah mengapa latensi kami tetap sub-detik (<25 ms) tanpa kehilangan keunggulan deteksi graf relasional."),
        ("Q2: Bagaimana jika ada nasabah baru yang belum ada di graf relasi GNN? Apakah transaksi langsung lolos?",
         "Sama sekali tidak. Formula fusi hibrida kami menerapkan mekanisme final_score = max(hybrid_score, rule_score). Rule Engine 13 indikator bertindak sebagai safety floor signal. Jika ada akun baru yang saldo awalnya tiba-tiba dikuras habis (drain-to-zero) atau bertransaksi di jam ganjil (02:00 WIB) dengan IP proxy VPN, Rule Engine akan langsung mengunci skor tinggi (hingga 100/BLOCK) tanpa bergantung pada skor GNN."),
        ("Q3: Bagaimana sistem Anda menjamin kepatuhan terhadap UU Pelindungan Data Pribadi (UU PDP No. 27/2022)?",
         "Kami menerapkan prinsip Privacy by Design. Di Forensic Dashboard, terdapat fitur Privacy Masking yang secara default menyamarkan Nama Nasabah, Nomor Rekening (****7890), dan NIK. Akses unmasking dibatasi ketat oleh RBAC dan setiap pembukaan data dicatat permanen di AuditLog mencakup identitas aktor, alasan, IP address, dan timestamp."),
        ("Q4: Apakah sistem Anda langsung memblokir rekening nasabah secara sepihak di bank nyata?",
         "Di lingkungan produksi perbankan, Crypto-Sentinel bertindak sebagai Decision Support System. Sistem melakukan penahanan sementara transaksi (Smart Circuit Breaker) pada kanal digital, namun pemblokiran permanen rekening nasabah dan pelaporan resmi ke PPATK tetap berada di bawah otorisasi manual Pejabat Kepatuhan (Compliance Officer / MLRO) berlisensi."),
        ("Q5: Apa bedanya solusi Anda dengan modul FDS bawaan Core Banking yang sudah dimiliki BPD saat ini?",
         "Modul bawaan Core Banking BPD umumnya hanya berbasis rule statis linier (misal limit transfer harian). Sindikat kejahatan masa kini mengakali hal ini dengan teknik smurfing bernominal kecil ke puluhan rekening mule. Modul lama tidak memiliki kemampuan melihat korelasi graf multi-hop. Crypto-Sentinel menambahkan lapisan intelijen graf relasional dan machine learning di depan core banking tanpa perlu mengganti sistem core banking lama bank.")
    ]
    for q, a in qa_list:
        add_callout(doc, q, a, bg="F8FAFC", border_color="0F172A")

    doc.add_page_break()

    # 7. SOP PRESENTASI FINAL & DEMO FLOW (HALAMAN 16-21)
    add_heading_styled(doc, "SOP Eksekusi Presentasi Final & Pembagian Peran Tim (Halaman 16–21)", level=1)
    
    add_heading_styled(doc, "1. Pembagian Peran Tim Saat Sesi Presentasi (Halaman 20)", level=2)
    role_data = [
        ("Lead Presenter", "Rifki Firmansyah", "Membawakan narasi pembuka (Slide 1-5), urgensi masalah, akar masalah, dampak terukur (Slide 9), dan kesimpulan penutup (Slide 12)."),
        ("Technical Presenter & Demo Operator", "Aam Setiana / Billy Jonathan", "Menjelaskan status kesiapan produk (Slide 6), arsitektur AI (Slide 7), benchmark performa (Slide 8), serta mengoperasikan Live Demo Prototype (60-90 detik)."),
        ("Business & Market Representative", "Billy Jonathan / Desta Erlangga", "Menjelaskan hasil validasi offtaker Bank bjb & Bank Kuningan (Slide 10), opsi adopsi Mode A/B, serta roadmap pilot 3 bulan (Slide 11)."),
        ("Q&A Domain Specialists", "Seluruh Anggota Sesuai Ownership", "AI & Algoritma: Rifki  |  Security & Core Banking: Billy  |  Dashboard UI & GNN: Aam  |  Backend & PPATK: Desta")
    ]
    t_role = doc.add_table(rows=len(role_data)+1, cols=3)
    t_role.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t_role)
    r_headers = ["Peran Sesi", "Anggota Tim", "Tanggung Jawab Spesifik"]
    r_w = [Inches(1.8), Inches(1.8), Inches(3.4)]
    for j, h in enumerate(r_headers):
        cell = t_role.cell(0, j)
        cell.width = r_w[j]
        set_cell_background(cell, "0F172A")
        set_cell_margins(cell, top=80, bottom=80, left=80, right=80)
        p = cell.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(role_data):
        for j, val in enumerate(row):
            cell = t_role.cell(i+1, j)
            cell.width = r_w[j]
            bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            p = cell.paragraphs[0]
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            if j <= 1:
                r.bold = True

    add_heading_styled(doc, "2. Skenario Live Demo Prototype Inti (Alur 5 Layar Inti — Durasi 60–90 Detik)", level=2)
    demo_steps = [
        ("Layar 1 — Executive Control Tower (Detik 00–15)", "Tunjukkan indikator service hijau, telemetri transaksi real-time, dan nilai dana yang berhasil diselamatkan (Blocked Value)."),
        ("Layar 2 — Live Monitoring & Intersepsi Mobile (Detik 16–35)", "Operator memicu transfer Rp 5.000.000 ke rekening bursa kripto di HP Android. Dalam <25 ms, muncul pop-up: 'Transaksi Tidak Dapat Diproses', saldo aman 100%, dan baris baru muncul di dashboard berstatus BLOCKED (Score 100)."),
        ("Layar 3 — Alert Detail & Risk Breakdown (Detik 36–55)", "Buka detail alert, tunjukkan dekomposisi skor: Rule Signal (Odd-hour), ML Feature Drain, dan GNN Mule Cluster. Tunjukkan faktor kontribusi SHAP."),
        ("Layar 4 — GNN Investigation Subgraph (Detik 56–75)", "Klik 'Telusuri Subgraf GNNExplainer'. Tunjukkan simpul pengirim, rekening mule, dan bursa kripto 3-hop menyala terang memisahkan sindikat dari nasabah normal."),
        ("Layar 5 — Compliance Action & Draft LTKM (Detik 76–90)", "Ubah status kasus menjadi RESOLVED, catat alasan investigasi di audit log, aktifkan saklar Privacy Masking (UU PDP), dan klik 'Terbitkan Draf LTKM PPATK' (dokumen goAML terbit dalam 3 detik). Selesai.")
    ]
    for st, sd in demo_steps:
        p_d = doc.add_paragraph(style='List Bullet')
        p_d.paragraph_format.space_before = Pt(1)
        p_d.paragraph_format.space_after = Pt(2)
        r_st = p_d.add_run(st + ": ")
        r_st.bold = True
        r_st.font.size = Pt(9)
        p_d.add_run(sd).font.size = Pt(9)

    add_heading_styled(doc, "3. Standar Pembuktian Kepada Dewan Juri (12-Item Evidence Pack Checklist)", level=2)
    evidence_items = [
        ("1. Live & Recorded Demo", "Skenario 1 transaction ID end-to-end (Mobile -> Gateway -> Dashboard)."),
        ("2. OpenAPI / Swagger Spec", "Kontrak API resmi Sentinel (:8000/docs) dan Expresso (:8080/docs)."),
        ("3. Architecture & Data-Flow Diagram", "Diagram 4-layer dan alur intersepsi pre-auth / post-auth."),
        ("4. Hasil Pengujian Jujur", "Matriks 8 Test Cases validasi stakeholder (3 PASS & 5 FAIL sebelum retest)."),
        ("5. Benchmark Latensi Resmi", "Evaluasi 308K PaySim (Mean: 5.67 ms, p95: 9.05 ms, p99: 12.23 ms)."),
        ("6. Model Card & Dataset Statement", "PaySim 308K + 12K edge cases lokal Indonesia (Bansos, SPP, QRIS, Kripto)."),
        ("7. Screenshot Source Badges", "Bukti visual LIVE · Sentinel API, LIVE · Core Banking API, dan DEMO FIXTURE."),
        ("8. Stakeholder Validation Notes", "Notulensi resmi uji keselarasan solusi bersama Bank bjb dan Bank Kuningan."),
        ("9. Deployment README & Bank Integration Kit", "Panduan instalasi 30–60 menit di docs/BANK_INTEGRATION_KIT.md."),
        ("10. Risk Register & Roadmap Pilot", "Matriks mitigasi 5 risiko dan rencana eksekusi pilot 3 bulan."),
        ("11. Sample Draft LTKM PPATK", "Draf resmi format goAML berlabel 'Draft/Synthetic' siap cetak PDF."),
        ("12. Kepatuhan Regulasi", "Pemetaan pasal POJK No. 8/2023, UU TPPU No. 8/2010, dan UU PDP No. 27/2022.")
    ]
    for e_title, e_desc in evidence_items:
        p_e = doc.add_paragraph(style='List Bullet')
        p_e.paragraph_format.space_before = Pt(1)
        p_e.paragraph_format.space_after = Pt(2)
        r_et = p_e.add_run(e_title + ": ")
        r_et.bold = True
        r_et.font.size = Pt(9)
        p_e.add_run(e_desc).font.size = Pt(9)

    add_heading_styled(doc, "4. Protokol Cadangan Bertingkat (3-Tier Backup Protocol)", level=2)
    p_proto = doc.add_paragraph()
    p_proto.add_run("Tier 1: Live Demo (Ponsel Android fisik via USB + Browser Dashboard lokal)\n"
                    "Tier 2: Recorded Demo MP4 (Video demo 75 detik, resolusi 1080p, tanpa audio promosi)\n"
                    "Tier 3: Screenshot Walkthrough Flow (Slide 6 & 7 di slide presentasi PowerPoint)")

    doc.add_page_break()

    # 8. SUPPORTING / APPENDIX SLIDES (HALAMAN 29-30)
    add_heading_styled(doc, "Supporting / Appendix Slides Master (Halaman 29–30)", level=1)
    p_ap_intro = doc.add_paragraph()
    p_ap_intro.add_run("Slide-slide pendukung berikut disimpan di belakang Slide 12 (Slide 13 s/d 22) untuk digunakan saat sesi Q&A mendalam:")
    
    appendices = [
        ("APPENDIX A: Detail Arsitektur Sistem & Spesifikasi Antarmuka (API & Data Flow)",
         "Progress Report: Bab 2.2 Integrasi Endpoint API  |  Blueprint: Section 01 Arsitektur & Section 04 API Spec",
         [
             ("Alur Transaksi Pre-Auth", "Mobile Banking -> Expresso CBS Gateway (:8080) -> Internal POST /api/v1/sentinel/evaluate -> AI Engine (:8000) -> Decision (<25ms)."),
             ("Spesifikasi Header SNAP BI", "X-TIMESTAMP (toleransi 5 menit), X-PARTNER-ID (KNG-BANK-001 / BJB-BANK-002), X-SIGNATURE (HMAC-SHA256 client_secret payload anti-tampering)."),
             ("Skema Database Relasional", "Tabel accounts (saldo, limit, is_blocked), transactions (log mutasi, FDS score), audit_logs (actor, role, action, reason, IP, timestamp), case_investigations (status OPEN/IN_REVIEW/ESCALATED/RESOLVED).")
         ]),
        ("APPENDIX B: Rekayasa Data & Matriks 29 Fitur Tabular (Dataset Deep-Dive)",
         "Progress Report: Bab 3.1 Dataset & 29 Fitur  |  Blueprint: Section 05 AI Dev Stage 1 & 2",
         [
             ("Komposisi PaySim Augmented", "320.606 baris total: 308.213 baris PaySim asli + 12.393 baris kasus lokal (3.200 Bansos, 2.800 SPP pendidikan, 2.500 QRIS UMKM, 3.893 pelarian kripto)."),
             ("29 Fitur Tabular", "Nominal (amount, balances), Rasio (balance_drain_ratio, error_balance), Temporal (hour_of_day, is_weekend, dormant_days, velocity_1h), One-hot (type_TRANSFER, purpose_GOVT, purpose_EDUC)."),
             ("Penyeimbangan Kelas", "SMOTE (Synthetic Minority Over-sampling) menyeimbangkan kelas fraud dari 3.31% menjadi 50:50 pada fase fitting 100 pohon keputusan.")
         ]),
        ("APPENDIX C: GraphSAGE GNN Deep-Dive (Topologi, Training & Formula XAI)",
         "Progress Report: Bab 3.2 GraphSAGE Training & GNNExplainer  |  Blueprint: Section 05 AI Dev Stage 3 & 4",
         [
             ("Statistik Graf Transaksi", "562.239 node rekening unik, 308.213 edge relasi transfer, 8-dimensi fitur node (degree, flow amount, fraud neighbor ratio)."),
             ("Training di GPU T4 (Colab)", "Optimizer Adam (lr=0.005), Class Weight Fraud=67.5. Epoch 1: Loss 0.5147, Val AUC 0.9995 -> Epoch 15: Loss 0.0004, Val AUC: 1.0000."),
             ("Formulasi GNNExplainer XAI", "max MI(Y, (G_s, F)) = H(Y) - H(Y | G=G_s, X=X*F). Memaksimalkan Mutual Information untuk memisahkan subgraf sindikat penting dari transaksi normal.")
         ]),
        ("APPENDIX D: Logika Lengkap 13 Indikator Rule Engine & Formula Haversine",
         "Progress Report: Bab 4.1 Pengujian Rule Engine  |  Blueprint: Section 03 Indikator 4 Signal Groups",
         [
             ("13 Indikator Pembobotan", "Channel Risiko (+25), Drain to Zero (+35), Odd-Hour (+25), Dormant (+30), Device Anomaly (+20), VPN Proxy (+20), Impossible Travel (+25), Threat Blacklist (+70), Smurfing (+45), Contextual Whitelist (-30)."),
             ("Formula Haversine Impossible Travel", "d = 2R * arcsin(sqrt(sin^2(d_lat/2) + cos(lat1)*cos(lat2)*sin^2(d_lon/2))). Kecepatan = d / delta_t. Jika kecepatan > 800 km/jam, memicu alert impossible travel.")
         ]),
        ("APPENDIX E: Estimasi Pasar (Market Sizing) & Model Keberlanjutan Finansial",
         "Progress Report: Bab 6.2 Proyeksi Kelayakan Finansial & TCO  |  Blueprint: Section 07 Roadmap & Skema Biaya",
         [
             ("Market Sizing Indonesia", "TAM: Rp 3,2 Triliun (107 Bank Umum & 1.400+ BPR); SAM: Rp 450 Miliar (26 BPD & BPR digital); SOM: Rp 25 Miliar (Adopsi awal BPR & BPD Jabar-Banten)."),
             ("Model Harga & Biaya", "Tier BPR: Rp 5 Juta/bln (hingga 50K transaksi + LTKM unlimited); Tier BPD: Rp 25 Juta/bln (multi-cabang + dedicated GNN + APOLO XML)."),
             ("Penghematan TCO Bank", "70% lebih terjangkau dibanding FDS enterprise multinasional dengan ROI pencegahan kerugian ganti rugi nasabah.")
         ]),
        ("APPENDIX F: Notulensi Lengkap 8 Test Cases Uji Lapangan (Solution Alignment)",
         "Progress Report: Bab 5.2 Evaluasi 8 Skenario Lapangan  |  Solution Alignment Notes: docs/solution_alignment_notes.md",
         [
             ("Bank bjb (TC-BJB-01 s/d 04)", "SOP AML (PASS); Anonimisasi PII UU PDP (Awalnya FAIL -> RESOLVED dengan Privacy Masking); Format goAML (PASS); Kalibrasi False Positive (Awalnya FAIL -> RESOLVED skor 60-84 masuk antrean REVIEW)."),
             ("Bank Kuningan (TC-KNG-01 s/d 04)", "Skenario Transaksi BPR (PASS); Impossible Travel (Awalnya FAIL -> RESOLVED validasi Haversine); Case Management (Awalnya FAIL -> RESOLVED tabel CaseInvestigation & AuditLog); Device Anomaly (Awalnya FAIL -> RESOLVED verifikasi device_id).")
         ]),
        ("APPENDIX G: Kerangka Kepatuhan Regulasi & Pemetaan Hukum (Regulatory Mapping)",
         "Progress Report: Bab 1.3 Kepatuhan Regulasi POJK & UU PDP  |  Blueprint: Section 00 Overview & Section 04 API Spec",
         [
             ("POJK No. 8/2023 (Strategi Anti-Fraud)", "Pilar 1 Pencegahan (Smart Circuit Breaker), Pilar 2 Deteksi (Multi-model AI), Pilar 3 Investigasi (CMS & Audit Trail), Pilar 4 Evaluasi (Kalibrasi Threshold Dinamis)."),
             ("UU No. 8/2010 (TPPU) & UU No. 27/2022 (PDP)", "Pasal 23 kewajiban LTKM dipangkas dari 3 hari ke 3 detik; Perlindungan PII nasabah melalui enkripsi dan pseudonimisasi default."),
             ("Standar APOLO OJK", "Formulir kepatuhan pengawas OJK dengan generator file ekspor XML terstandarisasi.")
         ]),
        ("APPENDIX H: Manajemen Risiko & Rencana Mitigasi (Risk Register)",
         "Progress Report: Bab 4.2 Analisis Keterbatasan & Failure Mitigations  |  Blueprint: Section 01 Arsitektur Fail-Closed",
         [
             ("Latensi API", "Mitigasi: Runtime inference menggunakan NumPy/Scikit lookup ringan tanpa GPU; timeout 50ms fallback."),
             ("False Positive Hari Raya/Gajian", "Mitigasi: Contextual Whitelist instansi resmi (-30 offset) dan antrean review manual 60-84."),
             ("Kebocoran Data PII di Layar", "Mitigasi: Privacy Masking aktif default; unmasking wajib otorisasi peran tinggi dan tercatat di immutable log."),
             ("Node Baru di Luar Graf", "Mitigasi: Formula max(hybrid, rule_score) memastikan Rule Engine mengunci fraud secara deterministik."),
             ("Kegagalan Database/Backend", "Mitigasi: Desain fail-closed transparan di API (DATA_SOURCES: LIVE vs ERROR); isolasi sandbox demo.")
         ]),
        ("APPENDIX I: Gantt Chart & Alokasi Sumber Daya Pilot 3 Bulan",
         "Progress Report: Bab 6.1 Roadmap Fase Pilot 3 Bulan  |  Blueprint: Section 07 Roadmap Pelaksanaan 3 Bulan",
         [
             ("Bulan 1 (Sandbox UAT & Security)", "M1: MoU & DPA UU PDP; M2: Setup environment terisolasi & sertifikat SNAP BI; M3: Load test p95/p99 latency; M4: User training staf kepatuhan."),
             ("Bulan 2 (Shadow Deployment Mode B)", "M5: Pemasangan CDC read-only database feed; M6: Pemantauan transaksi paralel; M7: Kalibrasi ulang threshold; M8: Simulasi skenario mule ring."),
             ("Bulan 3 (Audit SKAI & Sign-Off)", "M9: Audit bersama SKAI bank mitra; M10: Perhitungan metrik penghematan dana (ROI); M11: Laporan Uji Coba untuk OJK; M12: Pilot Acceptance Sign-Off.")
         ]),
        ("APPENDIX J: Checklist Berkas & Kelengkapan Presentasi (File & Packaging Kit)",
         "Progress Report: Bab 2.1 Lingkungan Pengujian Sandbox  |  Blueprint: Hero Stats & Nav Menu",
         [
             ("Berkas Digital (Online Pitching)", "1. Final_Pitch_Deck.pptx  |  2. Final_Pitch_Deck.pdf (Backup)  |  3. Recorded_Demo.mp4 (75s, 1080p)  |  4. Supporting_Appendix.pptx  |  5. Demo Link."),
             ("Berkas Fisik (Offline Pitching)", "Cetak Draf LTKM PPATK Resmi format goAML, Notulensi Uji Coba Bank bjb & Bank Kuningan, One-Click Launcher START-ALL.bat pada laptop offline.")
         ])
    ]

    for ap_title, ap_ref, ap_items in appendices:
        add_heading_styled(doc, ap_title, level=2)
        p_ref = doc.add_paragraph()
        p_ref.paragraph_format.space_before = Pt(1)
        p_ref.paragraph_format.space_after = Pt(3)
        r_tag = p_ref.add_run("📌 Rujukan Teknis: ")
        r_tag.bold = True
        r_tag.font.size = Pt(8.5)
        r_tag.font.color.rgb = RGBColor(2, 132, 199)
        p_ref.add_run(ap_ref).font.size = Pt(8.5)
        
        for k, v in ap_items:
            p_it = doc.add_paragraph(style='List Bullet')
            p_it.paragraph_format.space_before = Pt(1)
            p_it.paragraph_format.space_after = Pt(2)
            r_k = p_it.add_run(k + ": ")
            r_k.bold = True
            r_k.font.size = Pt(9)
            p_it.add_run(v).font.size = Pt(9)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    # Save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated DOCX at: {output_path}")

if __name__ == "__main__":
    out_file = r"d:\Crypto-Sentinel 2026\docs\FINAL_PITCH_DECK_CONTENT_2026.docx"
    build_pitch_deck_docx(out_file)
