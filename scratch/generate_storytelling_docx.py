"""
Script untuk menghasilkan berkas Microsoft Word (.docx) resmi:
docs/PANDUAN_STORYTELLING_GNN_DAN_WORKBENCH_FDS.docx
"""
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=140, right=140):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CBD5E1", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:left w:val="none"/>\n'
            f'  <w:right w:val="none"/>\n'
            f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
            f'  <w:insideV w:val="none"/>\n'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(level=level)
    h.paragraph_format.keep_with_next = True
    r = h.add_run(text)
    if level == 1:
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        r.font.size = Pt(15)
        r.bold = True
        r.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        r.font.size = Pt(12)
        r.bold = True
        r.font.color.rgb = RGBColor(2, 132, 199)
    elif level == 3:
        h.paragraph_format.space_before = Pt(8)
        h.paragraph_format.space_after = Pt(2)
        r.font.size = Pt(10.5)
        r.bold = True
        r.font.color.rgb = RGBColor(51, 65, 85)
    return h

def add_callout(doc, title, text, bg="F8FAFC", border_color="0284C7"):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, bg)
    set_cell_margins(cell, top=100, bottom=100, left=140, right=140)
    
    tcPr = cell._element.get_or_add_tcPr()
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>\n'
        f'  <w:left w:val="single" w:sz="24" w:space="0" w:color="{border_color}"/>\n'
        f'  <w:top w:val="none"/>\n'
        f'  <w:right w:val="none"/>\n'
        f'  <w:bottom w:val="none"/>\n'
        f'</w:tcBorders>'
    )
    tcPr.append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    if title:
        r_t = p.add_run(title + "\n")
        r_t.bold = True
        r_t.font.size = Pt(9.5)
        r_t.font.color.rgb = RGBColor(15, 23, 42)
    r_body = p.add_run(text)
    r_body.font.size = Pt(9)
    r_body.font.color.rgb = RGBColor(51, 65, 85)

def build_docx():
    doc = docx.Document()
    
    # Page Setup
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Title Hero
    t_hero = doc.add_table(rows=1, cols=1)
    t_hero.alignment = WD_TABLE_ALIGNMENT.CENTER
    c_hero = t_hero.cell(0, 0)
    c_hero.width = Inches(6.8)
    set_cell_background(c_hero, "09132E")
    set_cell_margins(c_hero, top=200, bottom=200, left=180, right=180)
    
    p_b = c_hero.paragraphs[0]
    p_b.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_b = p_b.add_run("PIDI DIGDAYA 2026 · TEAM EXPRESSO S1251\n")
    r_b.font.size = Pt(9)
    r_b.bold = True
    r_b.font.color.rgb = RGBColor(217, 119, 6)
    
    r_t = p_b.add_run("PANDUAN STORYTELLING DEEP-TECH & WORKBENCH FDS\n")
    r_t.font.size = Pt(17)
    r_t.bold = True
    r_t.font.color.rgb = RGBColor(255, 255, 255)
    
    r_s = p_b.add_run("Cara Kerja GNN, 15 Metrik Risiko Fraud Hibrida, Skenario Kontras, dan Alur Investigasi\n")
    r_s.font.size = Pt(10)
    r_s.font.color.rgb = RGBColor(148, 163, 184)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # 1. Seni Storytelling
    add_heading_styled(doc, "1. Seni Storytelling: Cara Menjelaskan GNN kepada Dewan Juri", level=1)
    add_callout(doc, "🗣️ Analogi 'Detektif vs Polisi Lalu Lintas' (Wajib Digunakan Saat Pitching!):",
                "Dewan Juri yang terhormat, bayangkan sistem FDS konvensional saat ini seperti Polisi Lalu Lintas yang berdiri di lampu merah. Dia hanya melihat satu mobil yang lewat: 'Apakah mobil ini melaju di atas kecepatan 80 km/jam? Jika tidak, silakan jalan.'\n\n"
                "Sindikat kejahatan masa kini sangat cerdik. Mereka tidak membawa 1 truk kontainer berisi uang curian Rp 500 juta, melainkan menyewa 100 sepeda motor kecil (rekening mule) yang masing-masing hanya membawa Rp 4,9 juta. Di mata polisi lalu lintas, setiap motor tampak normal dan legal.\n\n"
                "Crypto-Sentinel bekerja seperti Detektif Satelit berbasis Graph Neural Network (GNN). Sistem kami tidak hanya melihat satu motor, melainkan memetakan seluruh jaringan jalanan: dari mana motor-motor itu berangkat, siapa yang memberi komando, dan ke gudang mana (bursa kripto) mereka berkumpul secara serentak. Itulah mengapa kami bisa menangkap sindikat yang tidak pernah bisa dilihat oleh sistem lama!",
                bg="F0F9FF", border_color="0284C7")

    # 2. Anatomi Cara Kerja GNN
    add_heading_styled(doc, "2. Anatomi Cara Kerja GraphSAGE (Hamilton et al., NIPS 2017)", level=1)
    gnn_steps = [
        ("Babak 1: Neighbor Sampling (Pengambilan Sampel Tetangga)", "Alih-alih memproses jutaan rekening sekaligus, GraphSAGE mengambil sampel acak tetangga terdekat dalam radius k-hop (k=1, 2, 3) berdasarkan aliran transfer, kesamaan device fingerprint, dan IP address."),
        ("Babak 2: Message Passing & Aggregation (Pengumpulan Pesan Relasi)", "Setiap simpul mengumpulkan representasi fitur dari tetangganya menggunakan fungsi agregasi mean/max-pooling: h_N(v) = AGGREGATE({h_u, u in N(v)}). Sinyal bahaya dari bursa kripto merambat mundur ke rekening sasaran."),
        ("Babak 3: Node State Update (Pembaruan Status & Representasi Vektor)", "Menggabungkan status asal simpul dengan sinyal agregasi lingkungan: h_v = sigma( W · [ h_v || h_N(v) ] ). Menghasilkan vektor embedding 32-dimensi yang merefleksikan posisi sosial rekening.")
    ]
    for st, sd in gnn_steps:
        p = doc.add_paragraph(style='List Bullet')
        r_st = p.add_run(st + ": ")
        r_st.bold = True
        p.add_run(sd)

    # 3. 15 Metrik Risiko Fraud
    add_heading_styled(doc, "3. Simulasi 15 Metrik dan Indikator Fraud Hibrida", level=1)
    p_m_intro = doc.add_paragraph()
    p_m_intro.add_run("Sistem mengevaluasi 15 sinyal risiko secara simultan menggunakan sliding time-windows gabungan aturan deterministik dan probabilitas GNN:")
    
    metrics_data = [
        ("1", "Velocity (Kecepatan Pengurasan)", "Konvensional", "Akumulasi pengeluaran > Rp 20 Juta dalam durasi < 60 detik.", "+35 Poin"),
        ("2", "IP & Device Entropy", "Konvensional", "Kemunculan >= 3 IP / Device ID berbeda pada 1 akun dalam 1 jam.", "+20 Poin"),
        ("3", "Structuring / Smurfing", "Konvensional", "Serial transfer Rp 4,9 Juta - Rp 4,99 Juta (tepat di bawah limit audit).", "+45 Poin"),
        ("4", "Drain-to-Zero Ratio", "Konvensional", "Rasio amount/balance >= 0.98 (saldo dikuras hingga tersisa Rp 0).", "+35 Poin"),
        ("5", "Odd-Hour Anomaly", "Konvensional", "Transfer digital dieksekusi pada jam tidur nasabah (01:00 - 04:30 WIB).", "+25 Poin"),
        ("6", "Dormant Awakening", "Konvensional", "Akun pasif > 180 hari tiba-tiba transfer jumbo dalam < 15 menit.", "+30 Poin"),
        ("7", "Impossible Travel (Haversine)", "Konvensional", "Jarak fisik via rumus sferis Haversine dengan kecepatan > 800 km/jam.", "+25 Poin"),
        ("8", "Card Testing (Micro-Probe)", "Konvensional", "Transfer mikro Rp 10.000 (uji validasi) diikuti transfer saldo maksimal.", "+30 Poin"),
        ("9", "Sliding-Window Frequency", "Konvensional", "Frekuensi transaksi sukses > 10 kali dalam jendela 60 menit terakhir.", "+25 Poin"),
        ("10", "Contextual Whitelist", "Konvensional", "Tujuan lembaga resmi (Bansos, SPP Universitas, Pajak Daerah).", "-30 Poin (Diskon)"),
        ("11", "Mule Account Proximity", "Relasional GNN", "Jarak kosinus embedding GraphSAGE ke klaster rekening penampung.", "GNN Score (0-100)"),
        ("12", "Cyclic Flow / Circular Trading", "Relasional GNN", "Topologi graf tertutup di mana uang berputar (A -> B -> C -> A).", "+40 Poin"),
        ("13", "Common Beneficiary Fan-In", "Relasional GNN", ">= 5 pengirim berbeda mentransfer serentak ke 1 rekening tujuan baru.", "+45 Poin"),
        ("14", "Transaction Mixing (VASP)", "Relasional GNN", "Aliran multi-hop berujung ke bursa aset kripto (Indodax/Tokocrypto).", "+30 Poin"),
        ("15", "Threat Intel Graph Match", "Relasional GNN", "Kecocokan simpul graf dengan database penipu OJK IASC / PPATK.", "+70 Poin (Auto-Block)")
    ]
    t_m = doc.add_table(rows=len(metrics_data)+1, cols=5)
    set_table_borders(t_m)
    m_headers = ["No", "Nama Metrik", "Kategori", "Logika Evaluasi", "Bobot Skor"]
    m_w = [Inches(0.4), Inches(2.0), Inches(1.2), Inches(2.4), Inches(0.8)]
    for j, h in enumerate(m_headers):
        c = t_m.cell(0, j)
        c.width = m_w[j]
        set_cell_background(c, "0F172A")
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.color.rgb = RGBColor(255, 255, 255)
    for i, row in enumerate(metrics_data):
        for j, val in enumerate(row):
            c = t_m.cell(i+1, j)
            c.width = m_w[j]
            bg = "F8FAFC" if i % 2 == 1 else "FFFFFF"
            set_cell_background(c, bg)
            p = c.paragraphs[0]
            r = p.add_run(val)
            if j <= 1:
                r.bold = True

    # 4. Skenario Visual Kontras
    add_heading_styled(doc, "4. Skenario Visual Kontras di Layar Analis", level=1)
    contrast_data = [
        ("Parameter", "🟢 Skenario A: Transaksi Normal (Gaji)", "🔴 Skenario B: Sindikat Smurfing (Mule)"),
        ("Status & Skor", "ALLOW (Risk Score: 12 / 100)", "BLOCKED (Risk Score: 100 / 100)"),
        ("Sinyal Rule Terpicu", "Nominal wajar (Rp 3,5 Juta), jam kerja (10:15 WIB), Contextual Whitelist Payroll (-30).", "Smurfing 5x Rp 4,9 Juta (+45), Odd-Hour 02:45 WIB (+25), Drain-to-Zero (+35), Impossible Travel (+25)."),
        ("Visualisasi Graf GNN", "Simpul tunggal terhubung ke payroll resmi. Topologi bintang hijau, tanpa klaster mule.", "Simpul menyala merah terang 3-hop! Terlihat 10 pengirim mentransfer ke mule lalu ke VASP Indodax/Tokocrypto."),
        ("Keputusan Sistem", "Lolos seketika (<5 ms), saldo terkirim.", "Smart Circuit Breaker memutus transaksi (6 ms), saldo aman 100%, draf goAML terbit 3 detik.")
    ]
    t_c = doc.add_table(rows=len(contrast_data), cols=3)
    set_table_borders(t_c)
    for i, row in enumerate(contrast_data):
        for j, val in enumerate(row):
            c = t_c.cell(i, j)
            if i == 0:
                set_cell_background(c, "0F172A")
                p = c.paragraphs[0]
                r = p.add_run(val)
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
            else:
                bg = "ECFDF5" if j == 1 else ("FEF2F2" if j == 2 else "F8FAFC")
                set_cell_background(c, bg)
                p = c.paragraphs[0]
                r = p.add_run(val)
                if j == 0:
                    r.bold = True

    # 5. Alur Workbench
    add_heading_styled(doc, "5. Alur End-to-End Investigasi Forensik di Dashboard FDS", level=1)
    wb_steps = [
        ("1. Antrean Triage Real-Time", "Peringatan masuk seketika dengan badge status baku: BLOCK (Merah), REVIEW (Kuning), ALLOW (Hijau) dilengkapi timer SLA investigasi (SLA 4 Jam untuk Block)."),
        ("2. Panel Transparansi XAI", "Membuka dekomposisi skor multi-model (Rule, RF, GNN), SHAP feature contribution bar, dan visualisasi subgraf interaktif 3-hop GNNExplainer."),
        ("3. Customer 360 Live Database", "Menarik data CRA aktual dari NeonDB PostgreSQL (PEP status, CDD/EDD, mule probability, pekerjaan, pendapatan, dan mutasi 30 hari)."),
        ("4. Aksi Mitigasi & Draf LTKM", "Analis mengubah status kasus menjadi RESOLVED, merekomendasikan blokir permanen, atau mengklik 'Terbitkan Draf LTKM PPATK' (dokumen goAML terbit dalam 3 detik).")
    ]
    for st, sd in wb_steps:
        p = doc.add_paragraph(style='List Bullet')
        r_st = p.add_run(st + ": ")
        r_st.bold = True
        p.add_run(sd)

    # 6. Daya Jual B2B
    add_heading_styled(doc, "6. Arsitektur Real-Time Sub-Detik (<25 ms) & Daya Jual Startup B2B", level=1)
    add_callout(doc, "💼 4 Proposisi Nilai B2B bagi Bank Pembangunan Daerah & BPR:",
                "1. Ultra-Low Latency (<25 ms): Rata-rata inferensi 5.67 ms pada CPU lokal tanpa cluster GPU mahal di runtime bank.\n"
                "2. Non-Intrusive Plug-and-Play: Terhubung via Bank Integration Kit (docs/BANK_INTEGRATION_KIT.md) tanpa membongkar source code Core Banking lama.\n"
                "3. Hemat TCO hingga 70%: Skema berlangganan SaaS/On-Premise bertingkat (BPR Rp 5 Juta/bln, BPD Rp 25 Juta/bln) jauh lebih terjangkau dibanding vendor global (SAS/Actimize).\n"
                "4. Kepatuhan Native Indonesia: Memenuhi 100% regulasi POJK No. 8/2023, standar goAML PPATK, dan UU PDP No. 27/2022.",
                bg="F0FDF4", border_color="16A34A")

    output_path = r"d:\Crypto-Sentinel 2026\docs\PANDUAN_STORYTELLING_GNN_DAN_WORKBENCH_FDS.docx"
    doc.save(output_path)
    print(f"Successfully generated DOCX at: {output_path}")

if __name__ == "__main__":
    build_docx()
