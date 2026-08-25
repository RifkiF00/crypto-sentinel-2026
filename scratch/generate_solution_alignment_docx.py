import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D1D5DB", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'  <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def set_box_borders(table, color="2563EB", sz="8", val="single"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:insideV w:val="none"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def build_docx(output_path):
    doc = docx.Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Styles
    # Primary colors
    NAVY = RGBColor(15, 23, 42)      # #0F172A
    BLUE = RGBColor(37, 99, 235)     # #2563EB
    DARK_GRAY = RGBColor(51, 65, 85) # #334155
    MUTED = RGBColor(100, 116, 139)  # #64748B
    GREEN = RGBColor(22, 163, 74)    # #16A34A
    RED = RGBColor(220, 38, 38)      # #DC2626
    AMBER = RGBColor(217, 119, 6)    # #D97706

    # Normal Style
    normal_style = doc.styles['Normal']
    normal_style.font.name = 'Calibri'
    normal_style.font.size = Pt(10.5)
    normal_style.font.color.rgb = DARK_GRAY

    # Document Header Box (Hero Banner)
    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_cell = header_table.cell(0, 0)
    header_cell.width = Inches(6.9)
    set_cell_background(header_cell, "0F172A")
    set_cell_margins(header_cell, top=200, bottom=200, left=240, right=240)
    
    p_pre = header_cell.paragraphs[0]
    p_pre.paragraph_format.space_before = Pt(0)
    p_pre.paragraph_format.space_after = Pt(2)
    run_pre = p_pre.add_run("PROGRAM PIDI DIGDAYA HACKATHON & INKUBASI 2026 — DOKUMEN LEVEL 3")
    run_pre.font.size = Pt(8.5)
    run_pre.font.bold = True
    run_pre.font.color.rgb = RGBColor(56, 189, 248) # Sky blue
    
    p_title = header_cell.add_paragraph()
    p_title.paragraph_format.space_before = Pt(2)
    p_title.paragraph_format.space_after = Pt(4)
    run_title = p_title.add_run("Notulensi Uji Keselarasan Solusi\n(Solution Alignment Testing)")
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)
    
    p_sub = header_cell.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(2)
    p_sub.paragraph_format.space_after = Pt(0)
    run_sub = p_sub.add_run("Crypto-Sentinel 2026  |  Tim EXPRESSO S1251  |  Tanggal Pengujian: 25 Agustus 2026")
    run_sub.font.size = Pt(9.5)
    run_sub.font.color.rgb = RGBColor(203, 213, 225)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # Helper function for Section Heading
    def add_section_heading(title_text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(title_text)
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = NAVY

    # Helper function for Sub-heading (Test Case Title)
    def add_tc_heading(tc_id, title_text, status_badge=""):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        
        run_id = p.add_run(f"[{tc_id}] ")
        run_id.font.size = Pt(11)
        run_id.font.bold = True
        run_id.font.color.rgb = BLUE

        run_title = p.add_run(title_text)
        run_title.font.size = Pt(11)
        run_title.font.bold = True
        run_title.font.color.rgb = NAVY

        if status_badge:
            run_badge = p.add_run(f"  •  {status_badge}")
            run_badge.font.size = Pt(9.5)
            run_badge.font.bold = True
            if "PASS" in status_badge:
                run_badge.font.color.rgb = GREEN
            elif "FAIL" in status_badge:
                run_badge.font.color.rgb = RED
            else:
                run_badge.font.color.rgb = AMBER

    # Helper function for Key-Value Table
    def add_kv_table(data_dict, border_color="E2E8F0"):
        tbl = doc.add_table(rows=len(data_dict), cols=2)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_box_borders(tbl, color="CBD5E1")
        
        row_idx = 0
        for k, v in data_dict.items():
            cell_k = tbl.cell(row_idx, 0)
            cell_v = tbl.cell(row_idx, 1)
            cell_k.width = Inches(2.0)
            cell_v.width = Inches(4.9)
            
            set_cell_margins(cell_k, top=60, bottom=60, left=100, right=100)
            set_cell_margins(cell_v, top=60, bottom=60, left=100, right=100)
            
            # Key styling
            set_cell_background(cell_k, "F8FAFC")
            pk = cell_k.paragraphs[0]
            pk.paragraph_format.space_before = Pt(0)
            pk.paragraph_format.space_after = Pt(0)
            rk = pk.add_run(k)
            rk.font.bold = True
            rk.font.size = Pt(9.5)
            rk.font.color.rgb = DARK_GRAY
            
            # Value styling
            pv = cell_v.paragraphs[0]
            pv.paragraph_format.space_before = Pt(0)
            pv.paragraph_format.space_after = Pt(0)
            
            # Highlight status in value if present
            if k == "Status":
                rv = pv.add_run(v)
                rv.font.size = Pt(9.5)
                rv.font.bold = True
                if "PASS" in v:
                    rv.font.color.rgb = GREEN
                    set_cell_background(cell_v, "DCFCE7")
                elif "FAIL" in v:
                    rv.font.color.rgb = RED
                    set_cell_background(cell_v, "FEE2E2")
                else:
                    rv.font.color.rgb = AMBER
                    set_cell_background(cell_v, "FEF3C7")
            elif k == "Perubahan Solusi":
                rv = pv.add_run(v)
                rv.font.size = Pt(9.5)
                rv.font.bold = True
                rv.font.color.rgb = BLUE
                set_cell_background(cell_v, "EFF6FF")
            else:
                rv = pv.add_run(v)
                rv.font.size = Pt(9.5)
                rv.font.color.rgb = DARK_GRAY
                
            row_idx += 1
        
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # 1. INFORMASI UMUM
    add_section_heading("1. Informasi Umum Pengujian")
    info_dict = {
        "Nama Produk": "Crypto-Sentinel 2026",
        "Versi Inovasi": "v0.5.0 (SHAP Explainability + GNN GraphSAGE Hybrid)",
        "Tim Pengembang": "EXPRESSO S1251 — Rifki Firmansyah, Aam Setiana, Desta Erlangga, Billy Jonathan",
        "Offtaker Mitra 1": "Bank BJB (Bank Pembangunan Daerah Jawa Barat dan Banten) — Perusahaan BUMD",
        "Offtaker Mitra 2": "Bank Kuningan (BPR Kuningan) — Perusahaan BUMD",
        "Metode Pengujian": "Demo langsung + Diskusi penyelarasan kebutuhan + Uji coba sandbox simulasi",
        "Lingkungan Uji": "Sandbox Lokal (AI Engine :8000, Core Banking :8080, Dashboard :5173)"
    }
    add_kv_table(info_dict)

    # 2. BAGIAN A — BANK BJB
    add_section_heading("2. Bagian A — Pengujian & Notulensi Bersama Bank BJB")
    
    # TC-BJB-01
    add_tc_heading("TC-BJB-01", "Penyelarasan Fitur Sistem dengan Kebutuhan AML / APU-PPT", "PASS (dengan catatan)")
    add_kv_table({
        "ID Test Case": "TC-BJB-01",
        "Nama Pengujian": "Penyelarasan fitur Crypto-Sentinel dengan kebutuhan AML/APU-PPT BJB",
        "Tujuan": "Memverifikasi apakah fitur utama sistem relevan dengan proses kepatuhan AML yang dijalankan BJB.",
        "Pelaksana": "Tim EXPRESSO S1251 + Tim IT & Compliance Bank BJB",
        "Langkah Uji": "1. Presentasi komponen: FDS AI, Rule Engine, Dashboard, LTKM Generator, Mobile App.\n2. Demo transaksi simulasi interaktif.\n3. Diskusi keselarasan dengan SOP unit kepatuhan APU-PPT BJB.",
        "Hasil Diharapkan": "Seluruh komponen utama dinilai relevan dengan alur deteksi transaksi mencurigakan BJB.",
        "Hasil Aktual": "Komponen relevan. Catatan: FDS internal BJB memiliki ratusan jenis alert vs 13 indikator Crypto-Sentinel — gap diakui dan memerlukan roadmap bertahap.",
        "Status": "PASS (dengan catatan)",
        "Tindak Lanjut": "Menyusun roadmap penambahan sub-indikator deteksi bertahap selama masa inkubasi."
    })

    # TC-BJB-02
    add_tc_heading("TC-BJB-02", "Uji Anonimisasi Data Nasabah di Dashboard", "FAIL → Perlu Penyesuaian")
    add_kv_table({
        "ID Test Case": "TC-BJB-02",
        "Nama Pengujian": "Anonimisasi data nasabah sesuai standar keamanan perbankan BJB & UU PDP",
        "Tujuan": "Memverifikasi apakah data sensitif nasabah terlindungi pada tampilan visual dashboard compliance.",
        "Pelaksana": "Tim EXPRESSO S1251 + Tim Compliance Bank BJB",
        "Langkah Uji": "1. Buka dashboard monitoring compliance.\n2. Periksa tampilan nomor rekening, nama lengkap, dan NIK di tabel transaksi.\n3. Evaluasi kepatuhan terhadap UU Perlindungan Data Pribadi (UU PDP No. 27/2022).",
        "Hasil Diharapkan": "Nomor rekening dan identitas nasabah disamarkan/dianonimkan (contoh: ****7890).",
        "Hasil Aktual": "FAIL — Data nasabah masih tampil dalam bentuk mentah (plain text) tanpa masking.",
        "Status": "FAIL",
        "Perubahan Solusi": "Implementasi pseudonimisasi/tokenisasi: Rekening → ****7890, Nama → Inisial Terenkripsi, NIK → ****XXXX pada seluruh layer UI.",
        "Tindak Lanjut": "Menerapkan fungsi masking data nasabah pada komponen TransactionTable dan AlertDetail di dashboard."
    })

    # TC-BJB-03
    add_tc_heading("TC-BJB-03", "Uji Format Dokumen Laporan LTKM / STR Generator", "PASS")
    add_kv_table({
        "ID Test Case": "TC-BJB-03",
        "Nama Pengujian": "Validasi format laporan LTKM terhadap standar pelaporan PPATK goAML",
        "Tujuan": "Memverifikasi apakah draf dokumen LTKM otomatis sesuai format regulasi perbankan Indonesia.",
        "Pelaksana": "Tim EXPRESSO S1251 + Tim Compliance Bank BJB",
        "Langkah Uji": "1. Picu transaksi risiko tinggi (skor >=85).\n2. Buka endpoint /str/generate dan unduh draf PDF/HTML.\n3. Bandingkan struktur dokumen dengan formulir standar PPATK goAML (UU No. 8/2010).",
        "Hasil Diharapkan": "Dokumen memuat identitas PJK, identitas nasabah terlapor, uraian transaksi mencurigakan, dan tanda tangan pejabat kepatuhan.",
        "Hasil Aktual": "PASS — Format dinilai sudah sangat baik, formal, dan mencerminkan standar pelaporan perbankan.",
        "Status": "PASS",
        "Tindak Lanjut": "Fitur dipertahankan tanpa perubahan mendasar; integrasi export PDF dioptimalkan."
    })

    # TC-BJB-04
    add_tc_heading("TC-BJB-04", "Uji Smart Circuit Breaker — Mitigasi Risiko False Positive", "FAIL → Perlu Penyesuaian")
    add_kv_table({
        "ID Test Case": "TC-BJB-04",
        "Nama Pengujian": "Evaluasi risiko false positive pada mekanisme pemblokiran otomatis real-time",
        "Tujuan": "Memastikan sistem tidak memblokir transaksi nasabah sah secara sepihak yang dapat merusak reputasi bank.",
        "Pelaksana": "Tim EXPRESSO S1251 + Tim IT & Compliance Bank BJB",
        "Langkah Uji": "1. Simulasikan transaksi legal nominal besar dengan perilaku anomali ringan (skor 60-84).\n2. Amati respons sistem (ALLOW / REVIEW / BLOCK).\n3. Evaluasi dampak operasional jika transaksi tersebut diblokir otomatis.",
        "Hasil Diharapkan": "Sistem memisahkan zona REVIEW (antrean verifikasi) dan auto-BLOCK (hanya skor ekstrim >=85).",
        "Hasil Aktual": "FAIL — Model awal cenderung terlalu agresif memblokir transaksi pada skor menengah tanpa keterlibatan analis.",
        "Status": "FAIL",
        "Perubahan Solusi": "Penerapan Human-in-the-Loop: Skor 60-84 berstatus REVIEW (memerlukan persetujuan analis); Skor >=85 berstatus BLOCK. Ditambahkan tombol Override untuk analis kepatuhan.",
        "Tindak Lanjut": "Kalibrasi ulang ambang batas risiko di rule engine dan penyediaan endpoint override alert."
    })

    # 3. BAGIAN B — BANK KUNINGAN
    add_section_heading("3. Bagian B — Pengujian & Notulensi Bersama Bank Kuningan")

    # TC-KNG-01
    add_tc_heading("TC-KNG-01", "Uji Kesesuaian Skenario Transaksi Sandbox BPR Kuningan", "PASS (dengan catatan)")
    add_kv_table({
        "ID Test Case": "TC-KNG-01",
        "Nama Pengujian": "Validasi skenario transaksi simulasi dummy profile nasabah BPR",
        "Tujuan": "Memverifikasi kesesuaian simulasi sandbox dengan profil transaksi nasabah BPR daerah.",
        "Pelaksana": "Tim EXPRESSO S1251 + Staf IT Bank Kuningan",
        "Langkah Uji": "1. Siapkan data nasabah simulasi daerah Kuningan & Jawa Barat.\n2. Lakukan transaksi RTOL dan SKNBI melalui aplikasi mobile simulator.\n3. Pantau intersep scoring dan pencatatan log pada dashboard.",
        "Hasil Diharapkan": "Transaksi simulasi terproses mulus dan teranalisis secara real-time.",
        "Hasil Aktual": "PASS dengan catatan — Transaksi berhasil. Namun profil nasabah masih generik, belum mencerminkan segmentasi UMKM/nasabah mikro BPR Kuningan.",
        "Status": "PASS (dengan catatan)",
        "Perubahan Solusi": "Pembuatan dataset profil nasabah spesifik BPR Kuningan dengan pola transaksi komparatif (nominal wajar BPR, frekuensi regional).",
        "Tindak Lanjut": "Penyediaan 50+ akun dummy profil nasabah Jawa Barat untuk pengujian sandbox lanjutan."
    })

    # TC-KNG-02
    add_tc_heading("TC-KNG-02", "Uji Deteksi Perpindahan Dana Lintas Kota (Impossible Travel)", "FAIL → Perlu Penyesuaian")
    add_kv_table({
        "ID Test Case": "TC-KNG-02",
        "Nama Pengujian": "Deteksi anomali geolokasi transaksi lintas kota dalam selang waktu singkat",
        "Tujuan": "Mendeteksi sindikat kejahatan yang mengakses rekening dari kota berbeda dalam waktu tidak masuk akal.",
        "Pelaksana": "Tim EXPRESSO S1251 + Staf IT Bank Kuningan",
        "Langkah Uji": "1. Kirim transaksi pertama dari IP/koordinat Kuningan (Lat: -6.97, Long: 108.49).\n2. Dalam selang <10 menit, kirim transaksi kedua dari IP/koordinat Jakarta (Lat: -6.20, Long: 106.80).\n3. Amati apakah sistem membangkitkan indikator anomali kecepatan lokasi.",
        "Hasil Diharapkan": "Sistem mendeteksi Impossible Travel (>100 km dalam <30 menit) dan menaikkan skor risiko.",
        "Hasil Aktual": "FAIL — Rule engine belum mengkalkulasi kecepatan perpindahan jarak antar-transaksi berturut-turut.",
        "Status": "FAIL",
        "Perubahan Solusi": "Penambahan sub-rule Impossible Travel: Kalkulasi jarak Haversine antar-transaksi; jika kecepatan perpindahan >300 km/jam, tambahkan penalti risiko +25 poin.",
        "Tindak Lanjut": "Implementasi rule geo_velocity_anomaly di rule engine pada sprint pengembangan berikutnya."
    })

    # TC-KNG-03
    add_tc_heading("TC-KNG-03", "Uji Ketersediaan Case Management System (CMS) Investigasi", "FAIL → Perlu Penyesuaian")
    add_kv_table({
        "ID Test Case": "TC-KNG-03",
        "Nama Pengujian": "Evaluasi alur investigasi kasus dan status ticketing penanganan alert fraud",
        "Tujuan": "Menyediakan alur kerja investigasi berjenjang (workflow) bagi tim analis kepatuhan internal bank.",
        "Pelaksana": "Tim EXPRESSO S1251 + Tim Compliance Bank Kuningan",
        "Langkah Uji": "1. Picu alert transaksi mencurigakan.\n2. Coba ubah status kasus (OPEN -> IN_REVIEW -> CLOSED).\n3. Input catatan investigasi analis dan telaah riwayat perubahan audit.",
        "Hasil Diharapkan": "Tersedia modul CMS komprehensif dengan status tracking, investigasi note, dan riwayat petugas penelaah.",
        "Hasil Aktual": "FAIL — Sistem saat ini hanya menyediakan tombol aksi tunggal (Resolve) tanpa alur kasus bertahap.",
        "Status": "FAIL",
        "Perubahan Solusi": "Pembangunan modul Case Management System (CMS): Tabel database case_logs, endpoint update status & catatan investigasi, serta antarmuka investigasi terdedikasi di dashboard.",
        "Tindak Lanjut": "Mengembangkan modul Case Management System sebagai prioritas fungsional dashboard."
    })

    # TC-KNG-04
    add_tc_heading("TC-KNG-04", "Uji Audit Trail & Kebijakan 1 Akun 1 Perangkat (Device Binding)", "FAIL → Perlu Penyesuaian")
    add_kv_table({
        "ID Test Case": "TC-KNG-04",
        "Nama Pengujian": "Deteksi login multi-perangkat dan binding identitas perangkat nasabah",
        "Tujuan": "Mencegah pengambilalihan akun (Account Takeover) dengan membatasi sesi aktif hanya pada perangkat terdaftar.",
        "Pelaksana": "Tim EXPRESSO S1251 + Staf IT Bank Kuningan",
        "Langkah Uji": "1. Login akun nasabah dari Perangkat A (Android ID terdaftar).\n2. Lakukan login simultan akun yang sama dari Perangkat B.\n3. Periksa apakah sistem menolak login atau mengirimkan notifikasi keamanan darurat.",
        "Hasil Diharapkan": "Sistem mendeteksi ketidakcocokan device_id dan mewajibkan verifikasi ganda / memblokir sesi kedua.",
        "Hasil Aktual": "FAIL — Belum ada binding device_id pada mobile client; login dari perangkat berbeda masih lolos tanpa alert.",
        "Status": "FAIL",
        "Perubahan Solusi": "Implementasi Device Binding: Simpan device_id unik pada tabel akun. Trigger alert keamanan jika terdeteksi login dari device_id baru.",
        "Tindak Lanjut": "Menambahkan parameter device_id pada mobile banking API service dan tabel database akun."
    })

    # 4. RINGKASAN REKAPITULASI
    add_section_heading("4. Ringkasan Rekapitulasi Pengujian")
    
    summary_data = [
        ("ID", "Nama Pengujian", "Mitra Offtaker", "Hasil Uji", "Rencana Perubahan Solusi"),
        ("TC-BJB-01", "Penyelarasan Fitur AML/APU-PPT", "Bank BJB", "PASS (Catatan)", "Roadmap penambahan indikator bertahap"),
        ("TC-BJB-02", "Anonimisasi Data Nasabah", "Bank BJB", "FAIL", "Pseudonimisasi rekening & identitas nasabah di UI"),
        ("TC-BJB-03", "Format Laporan LTKM/STR", "Bank BJB", "PASS", "Format dipertahankan sesuai standar PPATK goAML"),
        ("TC-BJB-04", "Mitigasi False Positive Breaker", "Bank BJB", "FAIL", "Mekanisme Human-in-the-Loop & tombol override analis"),
        ("TC-KNG-01", "Skenario Sandbox Nasabah BPR", "Bank Kuningan", "PASS (Catatan)", "Dataset dummy profile lokal Jawa Barat"),
        ("TC-KNG-02", "Deteksi Lintas Kota (Geo)", "Bank Kuningan", "FAIL", "Rule Impossible Travel (geo_velocity_anomaly)"),
        ("TC-KNG-03", "Case Management System (CMS)", "Bank Kuningan", "FAIL", "Modul CMS: OPEN -> IN_REVIEW -> CLOSED + Note"),
        ("TC-KNG-04", "Device Binding 1 Akun 1 Device", "Bank Kuningan", "FAIL", "Pencatatan device_id & alert multi-device login")
    ]
    
    sum_tbl = doc.add_table(rows=len(summary_data), cols=5)
    sum_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_box_borders(sum_tbl, color="CBD5E1")
    
    col_widths = [Inches(0.9), Inches(1.8), Inches(1.1), Inches(1.1), Inches(2.0)]
    
    for r_idx, row in enumerate(summary_data):
        is_header = (r_idx == 0)
        for c_idx, val in enumerate(row):
            cell = sum_tbl.cell(r_idx, c_idx)
            cell.width = col_widths[c_idx]
            set_cell_margins(cell, top=60, bottom=60, left=60, right=60)
            
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            
            if is_header:
                set_cell_background(cell, "1E293B")
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                if c_idx == 3: # Hasil Uji column
                    run.font.bold = True
                    if "PASS" in val:
                        set_cell_background(cell, "DCFCE7")
                        run.font.color.rgb = GREEN
                    elif "FAIL" in val:
                        set_cell_background(cell, "FEE2E2")
                        run.font.color.rgb = RED
                else:
                    if r_idx % 2 == 1:
                        set_cell_background(cell, "F8FAFC")
                    run.font.color.rgb = DARK_GRAY

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # 5. PERUBAHAN YANG SUDAH SELESAI DIIMPLEMENTASI
    add_section_heading("5. Perubahan Solusi yang Telah Selesai Diimplementasi")
    
    impl_data = [
        ("Komponen Solusi", "Implementasi Teknis", "Respon Feedback Mitra"),
        ("Multi-Partner SNAP BI", "Autentikasi HMAC-SHA256 mendukung KNG-PARTNER & BJB-PARTNER secara simultan di expresso-api.", "Menjawab kebutuhan integrasi multi-bank (BJB & Kuningan) dengan secret key mandiri."),
        ("Koreksi Payload Mobile", "Penambahan field method (RTOL / SKNBI / OVERBOOKING) pada BankKuninganApiService client.", "Menyelaraskan alur pengiriman parameter transaksi core banking Bank Kuningan."),
        ("SHAP Explainability AI", "Integrasi SHAP TreeExplainer pada endpoint deteksi fraud untuk mengungkap kontribusi top-5 fitur.", "Menjawab kekhawatiran Bank BJB terkait transparansi alasan pemblokiran transaksi (Explainable AI).")
    ]
    
    impl_tbl = doc.add_table(rows=len(impl_data)+1, cols=3)
    impl_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_box_borders(impl_tbl, color="CBD5E1")
    
    headers_impl = ("Komponen Solusi", "Implementasi Teknis", "Respon Feedback Mitra")
    for c_idx, val in enumerate(headers_impl):
        cell = impl_tbl.cell(0, c_idx)
        set_cell_background(cell, "1E293B")
        set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(val)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(255, 255, 255)
        
    for r_idx, row in enumerate(impl_data):
        for c_idx, val in enumerate(row):
            cell = impl_tbl.cell(r_idx+1, c_idx)
            if (r_idx+1) % 2 == 1:
                set_cell_background(cell, "F8FAFC")
            set_cell_margins(cell, top=60, bottom=60, left=80, right=80)
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(val)
            run.font.size = Pt(8.5)
            if c_idx == 0:
                run.font.bold = True
                run.font.color.rgb = BLUE
            else:
                run.font.color.rgb = DARK_GRAY

    # Footer note
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    p_foot = doc.add_paragraph()
    p_foot.paragraph_format.space_before = Pt(8)
    p_foot.paragraph_format.space_after = Pt(0)
    run_foot = p_foot.add_run(
        "Dokumen notulensi uji keselarasan solusi ini disusun secara resmi oleh Tim EXPRESSO S1251 sebagai bukti pemenuhan "
        "kriteria Solution Alignment (Level 3) Program PIDI Digdaya Hackathon & Inkubasi 2026."
    )
    run_foot.font.size = Pt(8.5)
    run_foot.font.italic = True
    run_foot.font.color.rgb = MUTED

    doc.save(output_path)
    print(f"[OK] Document successfully generated at: {output_path}")

if __name__ == "__main__":
    out = r"d:\Crypto-Sentinel 2026\docs\solution_alignment_notes.docx"
    build_docx(out)
